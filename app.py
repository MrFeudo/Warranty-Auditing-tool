# -*- coding: utf-8 -*-
"""
Warranty Audit Assistant - versión interna bilingüe

Instalación:
    py -m pip install streamlit pandas openpyxl xlsxwriter google-genai python-docx

Ejecución:
    streamlit run app.py

Qué hace:
- Permite pegar una lista de claims en una tabla editable tipo Excel.
- Permite trabajar por bloques: primero documentación para todas las claims y después piezas viejas en almacén.
- Calcula automáticamente la puntuación:
    Claim document checklist I = 58 puntos
    Claim old parts checklist II = 42 puntos
    Total = 100 puntos
- Por defecto todos los apartados arrancan como OK / puntuación máxima.
- No aplica = puntuación máxima del apartado.
- Genera automáticamente el comentario general desde las observaciones por apartado, sin botón manual.
- Exporta boletín en español para dealer usando claim local CO...
- Exporta scorecard en inglés para HQ usando identificador HQ/IDMS/TAC/2810...
- Rellena la hoja/page 3 "Improvement of claim issues III" con las observaciones de los campos.
- En exportaciones EN, traduce las observaciones con un glosario local de auditoría.
- Permite registrar deducción parcial o total. El botón "No es garantía" pone la claim a 0/100 y marca deducción total.
"""

from __future__ import annotations

import json
import os
import re
import zipfile
import unicodedata
from dataclasses import dataclass
from datetime import datetime, date
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st


# =============================================================================
# CONFIGURACIÓN DE PUNTUACIÓN
# =============================================================================

@dataclass(frozen=True)
class AuditOption:
    key: str
    label_es: str
    label_en: str
    points: Optional[int]
    status: str


@dataclass(frozen=True)
class AuditCheck:
    block_key: str
    key: str
    label_es: str
    label_en: str
    max_points: int
    options: Tuple[AuditOption, ...]
    guidance_es: str
    guidance_en: str


PENDING = AuditOption("pending", "Pendiente de revisar", "Pending review", None, "Pendiente")


def options_0_5_7(max_points: int = 7) -> Tuple[AuditOption, ...]:
    return (
        PENDING,
        AuditOption("ok", "OK / conforme", "OK / compliant", max_points, "OK"),
        AuditOption("partial", "Parcial / incompleto", "Partial / incomplete", 5, "Parcial"),
        AuditOption("nok", "NOK / falta o no conforme", "NOK / missing or non-compliant", 0, "NOK"),
        AuditOption("na", "No aplica", "Not applicable", max_points, "N/A"),
    )


def options_0_6() -> Tuple[AuditOption, ...]:
    return (
        PENDING,
        AuditOption("ok", "OK / conforme", "OK / compliant", 6, "OK"),
        AuditOption("nok", "NOK / no conforme", "NOK / non-compliant", 0, "NOK"),
        AuditOption("na", "No aplica", "Not applicable", 6, "N/A"),
    )


def options_date_0_3_6() -> Tuple[AuditOption, ...]:
    return (
        PENDING,
        AuditOption("ok", "0 a 5 días laborables", "0 to 5 working days", 6, "OK"),
        AuditOption("partial", "5 a 15 días laborables", "5 to 15 working days", 3, "Parcial"),
        AuditOption("nok", "Más de 15 días laborables", "More than 15 working days", 0, "NOK"),
        AuditOption("na", "No aplica", "Not applicable", 6, "N/A"),
    )


def options_old_binary_0_7() -> Tuple[AuditOption, ...]:
    return (
        PENDING,
        AuditOption("ok", "OK / correcta", "OK / correct", 7, "OK"),
        AuditOption("nok", "NOK / incorrecta", "NOK / incorrect", 0, "NOK"),
        AuditOption("na", "No aplica", "Not applicable", 7, "N/A"),
    )



DOCUMENT_CHECKS: List[AuditCheck] = [
    AuditCheck(
        "document", "doc_or", "OR", "Repair order", 7, options_0_5_7(),
        "Comprobar si existe la orden de reparación/reclamación, si la documentación es completa y si la firma está correcta.",
        "Check whether the repair/warranty order exists, whether the relevant documentation is complete and whether the signature is correct.",
    ),
    AuditCheck(
        "document", "doc_parts_order", "Pedido piezas / albarán", "Parts order / delivery note", 7, options_0_5_7(),
        "Comprobar pedido/albarán de piezas y que la documentación sea completa, razonable y conforme.",
        "Check whether the parts order/delivery note exists and whether the documentation is complete, reasonable and compliant.",
    ),
    AuditCheck(
        "document", "doc_previous_or", "OR previa", "Previous repair order", 7, options_0_5_7(),
        "Solo para reclamaciones de repuestos o reparaciones anteriores. Si no procede, marcar No aplica.",
        "Only for parts claims or previous repairs. If not applicable, mark as Not applicable.",
    ),
    AuditCheck(
        "document", "doc_evidence", "Evidencias", "Evidence", 7, options_0_5_7(),
        "Comprobar si faltan evidencias adjuntas o si las evidencias no son correctas/suficientes.",
        "Check whether evidence is missing or whether the evidence provided is not correct/sufficient.",
    ),
    AuditCheck(
        "document", "doc_causal_part", "Pieza causa correcta", "Correct causal part", 6, options_0_6(),
        "Comprobar si la pieza principal dañada es correcta, razonable y coherente con la avería.",
        "Check whether the main damaged part is correct, reasonable and consistent with the failure.",
    ),
    AuditCheck(
        "document", "doc_labor", "Mano de obra", "Labour", 6, options_0_6(),
        "Comprobar si los tiempos de mano de obra son correctos, no repetitivos y ajustados al estándar.",
        "Check whether labour times are correct, not duplicated and aligned with the standard labour time.",
    ),
    AuditCheck(
        "document", "doc_aux_material", "Material auxiliar", "Auxiliary material", 6, options_0_6(),
        "Comprobar que el material auxiliar/consumible cumpla normativa, campo correcto y cantidad razonable.",
        "Check whether auxiliary/consumable material complies with the policy, is entered in the correct field and has a reasonable quantity.",
    ),
    AuditCheck(
        "document", "doc_dates", "Fecha/hora envío Claim", "Claim submission and repair date", 6, options_date_0_3_6(),
        "La claim debe enviarse dentro de plazo tras completar la reparación.",
        "The claim must be submitted within the allowed period after the repair is completed.",
    ),
    AuditCheck(
        "document", "doc_vin", "VIN", "VIN", 6, options_0_6(),
        "Comprobar que VIN, kilometraje, fecha de reparación, tipo de reclamación y datos del vehículo sean correctos.",
        "Check that VIN, mileage, repair date, claim type and vehicle data are correct.",
    ),
]

OLD_PARTS_CHECKS: List[AuditCheck] = [
    AuditCheck(
        "old_parts", "old_management", "Gestión piezas viejas", "Old parts management", 7, options_0_5_7(),
        "Las piezas viejas deben estar ordenadas, localizables y disponibles durante la auditoría.",
        "Old parts must be arranged, traceable and available during the audit.",
    ),
    AuditCheck(
        "old_parts", "old_label", "Etiquetado pieza vieja", "Old part labelling", 7, options_0_5_7(),
        "La etiqueta debe incluir datos básicos del vehículo, claim, referencia, causa del daño, etc.",
        "The label must include basic vehicle data, claim number, part number, damage cause, etc.",
    ),
    AuditCheck(
        "old_parts", "old_causal_part", "Pieza causa", "Causal part", 7, options_old_binary_0_7(),
        "Comprobar que la pieza antigua es la causa real, consistente con modelo/vehículo y fecha de producción.",
        "Check whether the old part is the real causal part and is consistent with the vehicle/model and production date.",
    ),
    AuditCheck(
        "old_parts", "old_failure_info", "Info tipo fallo pieza causa", "Causal part failure information", 7, options_old_binary_0_7(),
        "Comprobar que la información del fallo sea coherente con la pieza y las evidencias disponibles.",
        "Check whether the failure information is consistent with the part and with the available evidence.",
    ),
    AuditCheck(
        "old_parts", "old_destruction", "Destrucción pieza vieja", "Old part destruction", 7, options_0_5_7(),
        "Las piezas viejas deben destruirse siguiendo el proceso y sin posibilidad de reutilización.",
        "Old parts must be destroyed according to the process and must not be reusable.",
    ),
    AuditCheck(
        "old_parts", "old_destruction_certificate", "Certificado destrucción piezas viejas", "Old parts destruction certificate", 7, options_0_5_7(),
        "El certificado/informe de destrucción debe existir y archivarse en plazo.",
        "The destruction certificate/report must exist and be archived within the required period.",
    ),
]

ALL_SCORING_CHECKS = DOCUMENT_CHECKS + OLD_PARTS_CHECKS
ALL_CHECKS = ALL_SCORING_CHECKS
CHECK_BY_KEY = {check.key: check for check in ALL_CHECKS}
MAX_DOCUMENT_POINTS = sum(check.max_points for check in DOCUMENT_CHECKS)
MAX_OLD_PARTS_POINTS = sum(check.max_points for check in OLD_PARTS_CHECKS)
MAX_TOTAL_POINTS = MAX_DOCUMENT_POINTS + MAX_OLD_PARTS_POINTS

GEMINI_MODEL_DEFAULT = "gemini-2.5-flash"

ACTIVE_DEALERS: List[str] = [
    "ACAI MOTOR MÁLAGA", "ALFAVISA BILBAO", "ALIMOTOR ELCHE", "ANFERPA SEGOVIA",
    "AUTO YALDE CALAHORRA", "AUTO YALDE LOGROÑO", "AUTOCAM MOTOR VILAFRANCA",
    "AUTOCAM VILANOVA", "AUTOCYL PALENCIA", "AUTOCYL VALLADOLID", "AUTOVIDAL PALMA DE MALLORCA",
    "AVANTI GRANADA", "AXIS MOTORS", "BLENDIO LAREDO", "BLENDIO LUGO", "BLENDIO OURENSE",
    "BLENDIO OVIEDO", "BLENDIO SANTANDER", "BLENDIO TORRELAVEGA", "BORJAMOTOR ALICANTE",
    "CERVERA AVILA", "CERVERA SALAMANCA", "CHINARES GUADALAJARA", "DILOAUTOJAEN",
    "DUMOSA BENAVENTE", "ESLAUTO LEON", "FIMALAGA MÁLAGA", "FIMALAGA MARBELLA",
    "GRUP BASOLS IGUALADA", "GRUPO JULIAN BURGOS", "GRUPO NIETO MÁLAGA", "GRUPO NIETO MARBELLA",
    "HIMASA SEDAVÍ", "JEMOYA SORIA", "JOVERAUTO MELILLA", "LASACAR MIRANDA DE EBRO",
    "LASACAR VITORIA", "LEPAS AUTOCAM VILANOVA", "LEPAS AUTOVIVO SANT BOI",
    "LEPAS BASOLS IGUALADA", "LEPAS BASOLS VIC", "LEPAS GAMBOA MAJADAHONDA",
    "LEPAS JULIÁN BURGOS", "LEPAS MONECAR SAGUNTO", "LEPAS PREMIER VITORIA",
    "LEPAS RAFAEL AFONSO LAS PALMAS", "LEPAS RESNOVA CORUÑA", "LEPAS RESNOVA VIGO",
    "LEPAS TECNOTARRACO TARRAGONA", "LEPAS TUMASA HUESCA", "LEPAS VALLESCAR SABADELL",
    "LEPAS VALLESCAR TERRASSA", "LEPAS ZEN MOTOR GIPUZKOA", "LEPAS ZEN MOTOR ZARAGOZA",
    "M AUTOMOCIÓN ALCALÁ", "M AUTOMOCIÓN BCN (GRAN VÍA)", "M AUTOMOCIÓN BCN GUAYAQUIL",
    "M AUTOMOCIÓN CASTELLÓN", "M AUTOMOCIÓN GERONA", "M AUTOMOCIÓN MATARÓ",
    "M TECNIK ALCALÁ DE HENARES", "M TECNIK BARCELONA MAQUINISTA", "M TECNIK CASTELLÓN",
    "M TECNIK FIGUERES", "M TECNIK GERONA", "M TECNIK MATARÓ", "M TECNIK VINAROZ",
    "MARTIN LIZAGA TERUEL", "MAS AUTO LEGANÉS", "MAVEN BADAJOZ", "MAVEN CÁCERES",
    "MAVEN DON BENITO", "MAVEN MÉRIDA", "MAVEN PLASENCIA", "MOLL MOTOR DENIA",
    "MOLL MOTOR GANDIA", "MOLL VALENCIA", "MONECAR CUENCA", "MOTOR NACIENTE LEGANÉS",
    "MOVINSUR GRANADA", "MOVINSUR JAÉN", "MOVINSUR MOTRIL", "MY CARS CÓRDOBA",
    "NOVACAR BCN SANT BOI", "PALAUSA ZAMORA", "PROCHERY ALBACETE", "PROCHERY CARTAGENA",
    "PROCHERY MURCIA", "PRUNA CAR GO GRANOLLERS", "RAFAEL AFONSO AGUIMES",
    "RAFAEL AFONSO LANZAROTE", "RAFAEL AFONSO LAS PALMAS", "RAFAEL AFONSO TENERIFE",
    "RESNOVA MOTOR CORUÑA", "RESNOVA MOTOR GIJÓN", "RESNOVA MOTOR NARÓN",
    "RESNOVA MOTOR OVIEDO", "RESNOVA MOTOR SANTIAGO", "RESNOVA MOTOR VIGO", "SEGRE LLEIDA",
    "SEGRE MOTORS LERIDA", "SERTECAUTO PONFERRADA", "SYRSA ALGECIRAS", "SYRSA ALMERIA",
    "SYRSA EJIDO", "SYRSA HUELVA", "SYRSA SEVILLA", "TALAUTO CAZALEGAS", "TALAUTO TOLEDO",
    "TALLERES CHINARES", "TECNOTARRACO TARRAGONA", "TERRY MOBILITY JERÉZ",
    "TRADECAR GAMBOA ALCORCÓN", "TRADECAR GAMBOA MADRID", "TRADECAR GAMBOA MAJADAHONDA",
    "TRADECAR GAMBOA RIVAS", "TUMASA HUESCA", "TUMASA MONZÓN", "UNIONE ALCAZAR DE SAN JUAN",
    "UNIONE CIUDAD REAL", "VALLESCAR SABADELL", "VALLESCAR TERRASSA", "VIAN ALCORCÓN",
    "VIAN AUTOMOBILE VILLALBA", "VIAN MÓSTOLES", "VIAN NAVARRA", "ZEN MOTOR OLABERRIA",
    "ZEN MOTOR PAMPLONA", "ZEN MOTOR SAN SEBASTIÁN", "ZEN MOTOR ZARAGOZA",
]


# =============================================================================
# TEXTOS / TRADUCCIONES
# =============================================================================

TEXT = {
    "es": {
        "app_title": "Warranty Audit Assistant",
        "claim_id": "Claim dealer / España",
        "other_id": "Identificador HQ / IDMS",
        "scorecard": "Boletín de notas",
        "document_section": "Claim document checklist I",
        "old_parts_section": "Claim old parts checklist II",
        "improvement_sheet": "Improvement of claim issues III",
        "evaluation_content": "Evaluation content",
        "claim_no": "Claim No.",
        "local_claim": "Claim España / Dealer",
        "hq_claim": "Claim HQ / IDMS",
        "dealer": "Dealer",
        "vin": "VIN",
        "model": "Modelo",
        "amount": "Importe",
        "deduction_type": "Tipo deducción",
        "deduction_amount": "Deducción",
        "deduction_reason": "Motivo deducción",
        "no_deduction": "Sin deducción",
        "partial_deduction": "Deducción parcial",
        "full_deduction": "Deducción total",
        "not_warranty": "No es garantía",
        "doc_score": "Documentación /58",
        "old_score": "Piezas viejas /42",
        "total_score": "Total /100",
        "success": "% éxito",
        "result": "Resultado",
        "pending": "Pendientes",
        "comments": "Comentarios",
        "action_plan_title": "Plan de mejora",
        "auditable_list": "Lista de parámetros auditables",
        "exception_comments": "Excepciones/comentarios",
        "observations": "Observaciones",
        "countermeasure": "Contramedida",
        "no_deviations": "Sin desviaciones puntuables ni observaciones registradas",
        "generic_countermeasure": "Reforzar el cumplimiento del criterio y revisar la documentación antes del envío de la claim.",
        "excellent": "Excelente",
        "correct": "Correcto",
        "improvable": "Mejorable",
        "critical": "Crítico",
        "report_title": "Informe de auditoría",
        "executive_summary": "Resumen ejecutivo",
        "block_result": "Resultado por bloque",
        "improvement_areas": "Principales áreas de mejora",
        "lowest_claims": "Claims con menor puntuación",
        "conclusion": "Conclusión",
        "not_translated_note": "Nota: las observaciones manuales se incluyen tal como se han escrito.",
    },
    "en": {
        "app_title": "Warranty Audit Assistant",
        "claim_id": "HQ / IDMS claim ID",
        "other_id": "Dealer / Spanish claim ID",
        "scorecard": "Scorecard",
        "document_section": "Claim document checklist I",
        "old_parts_section": "Claim old parts checklist II",
        "improvement_sheet": "Improvement of claim issues III",
        "evaluation_content": "Evaluation content",
        "claim_no": "Claim No.",
        "local_claim": "Spanish / dealer claim",
        "hq_claim": "HQ / IDMS claim",
        "dealer": "Dealer",
        "vin": "VIN",
        "model": "Model",
        "amount": "Amount",
        "deduction_type": "Deduction type",
        "deduction_amount": "Deduction",
        "deduction_reason": "Deduction reason",
        "no_deduction": "No deduction",
        "partial_deduction": "Partial deduction",
        "full_deduction": "Full deduction",
        "not_warranty": "Not covered by warranty",
        "doc_score": "Document checklist /58",
        "old_score": "Old parts checklist /42",
        "total_score": "Total /100",
        "success": "Success %",
        "result": "Result",
        "pending": "Pending items",
        "comments": "Comments",
        "action_plan_title": "Improvement plan",
        "auditable_list": "Auditable parameter list",
        "exception_comments": "Exceptions/comments",
        "observations": "Observations",
        "countermeasure": "Countermeasure",
        "no_deviations": "No score deviations or observations recorded",
        "generic_countermeasure": "Reinforce compliance with the criterion and review the documentation before claim submission.",
        "excellent": "Excellent",
        "correct": "Correct",
        "improvable": "Needs improvement",
        "critical": "Critical",
        "report_title": "Audit report",
        "executive_summary": "Executive summary",
        "block_result": "Result by section",
        "improvement_areas": "Main improvement areas",
        "lowest_claims": "Lowest scoring claims",
        "conclusion": "Conclusion",
        "not_translated_note": "Note: manual observations are translated into English using the built-in audit glossary. Review uncommon wording before sending externally.",
    },
}

COUNTERMEASURES = {
    "doc_or": {
        "es": "Asegurar que la OR esté disponible, completa, firmada y trazable antes del envío de la claim.",
        "en": "Ensure the repair order is available, complete, signed and traceable before claim submission.",
    },
    "doc_parts_order": {
        "es": "Adjuntar y verificar el pedido de piezas/albarán correspondiente, asegurando coherencia con la reparación.",
        "en": "Attach and verify the corresponding parts order/delivery note, ensuring consistency with the repair.",
    },
    "doc_previous_or": {
        "es": "Verificar la reparación previa cuando aplique y adjuntar la OR/documentación soporte correspondiente.",
        "en": "Verify the previous repair when applicable and attach the related repair order/supporting documentation.",
    },
    "doc_evidence": {
        "es": "Reforzar la calidad de las evidencias adjuntas para demostrar diagnóstico, avería y reparación realizada.",
        "en": "Improve the quality of attached evidence to support diagnosis, failure and performed repair.",
    },
    "doc_causal_part": {
        "es": "Revisar que la pieza causa reclamada sea correcta y coherente con el diagnóstico y la reparación.",
        "en": "Check that the claimed causal part is correct and consistent with the diagnosis and repair.",
    },
    "doc_labor": {
        "es": "Ajustar la mano de obra al baremo y justificar cualquier tiempo adicional con desglose y evidencias.",
        "en": "Align labour with the standard time and justify any additional time with breakdown and evidence.",
    },
    "doc_aux_material": {
        "es": "Registrar el material auxiliar en el campo correcto, con desglose y justificación suficiente.",
        "en": "Enter auxiliary material in the correct field, with sufficient breakdown and justification.",
    },
    "doc_dates": {
        "es": "Controlar los plazos entre reparación y envío de la claim para evitar envíos fuera de plazo.",
        "en": "Monitor the period between repair completion and claim submission to avoid late submissions.",
    },
    "doc_vin": {
        "es": "Verificar VIN, kilometraje, fechas y datos del vehículo antes de enviar la claim.",
        "en": "Verify VIN, mileage, dates and vehicle data before claim submission.",
    },
    "old_management": {
        "es": "Mantener las piezas viejas ordenadas, localizables y disponibles durante el periodo de auditoría.",
        "en": "Keep old parts arranged, traceable and available during the audit period.",
    },
    "old_label": {
        "es": "Completar correctamente las etiquetas de piezas viejas con los datos requeridos de vehículo, claim y pieza.",
        "en": "Correctly complete old part labels with the required vehicle, claim and part data.",
    },
    "old_causal_part": {
        "es": "Conservar y presentar la pieza causa correcta, coherente con el vehículo y la avería reclamada.",
        "en": "Store and present the correct causal part, consistent with the vehicle and the claimed failure.",
    },
    "old_failure_info": {
        "es": "Documentar correctamente el tipo de fallo de la pieza causa y su coherencia con la evidencia.",
        "en": "Correctly document the causal part failure type and its consistency with the evidence.",
    },
    "old_destruction": {
        "es": "Seguir el proceso de destrucción de piezas viejas y evitar cualquier posibilidad de reutilización.",
        "en": "Follow the old parts destruction process and avoid any possibility of reuse.",
    },
    "old_destruction_certificate": {
        "es": "Archivar y presentar el certificado de destrucción dentro del periodo correspondiente.",
        "en": "Archive and provide the destruction certificate within the required period.",
    },
}


# =============================================================================
# UTILIDADES
# =============================================================================


def safe_str(value: Any) -> str:
    if value is None:
        return ""
    try:
        if isinstance(value, float) and pd.isna(value):
            return ""
    except Exception:
        pass
    return str(value).strip()


def json_default_serializer(value: Any):
    """Fallback para mandar datos a Gemini sin petar por tipos de pandas/numpy/datetime."""
    if isinstance(value, datetime):
        return value.isoformat(timespec="seconds")
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass
    try:
        # numpy int/float/bool, si aparecen al leer Excel o dataframes.
        if hasattr(value, "item"):
            return value.item()
    except Exception:
        pass
    return safe_str(value)


def normalize_text(value: Any) -> str:
    text = safe_str(value).lower()
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def sanitize_for_filename(value: Any, fallback: str = "item") -> str:
    text = safe_str(value) or fallback
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = re.sub(r"[^A-Za-z0-9_. -]+", "_", text).strip()
    text = re.sub(r"\s+", "_", text)
    return text or fallback


def build_audit_file_basename(dealer: str, auditor: str, when: Optional[datetime] = None) -> str:
    when = when or datetime.now()
    return f"{sanitize_for_filename(dealer, 'Dealer')}_{when.strftime('%Y%m%d')}_{sanitize_for_filename(auditor, 'Auditor')}"



def parse_audit_date(value: Any) -> date:
    """Devuelve una fecha válida para la auditoría."""
    if isinstance(value, date):
        return value
    text = safe_str(value)
    if text:
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                return datetime.strptime(text, fmt).date()
            except Exception:
                pass
    return datetime.now().date()


def audit_date_to_text(value: Any, language: str = "es") -> str:
    audit_date = parse_audit_date(value)
    if language == "en":
        return audit_date.strftime("%Y-%m-%d")
    return audit_date.strftime("%d/%m/%Y")


def build_audit_file_basename_from_date(dealer: str, auditor: str, audit_date_value: Any = None) -> str:
    audit_date = parse_audit_date(audit_date_value)
    return build_audit_file_basename(
        dealer,
        auditor,
        datetime.combine(audit_date, datetime.min.time()),
    )




def parse_amount_value(value: Any) -> float:
    """Convierte importes escritos como 1.234,56 €, 1234.56 o 1234 a float."""
    text = safe_str(value)
    if not text:
        return 0.0
    text = text.replace("€", "").replace(" ", "")
    text = re.sub(r"[^0-9,.-]", "", text)
    if not text:
        return 0.0
    # Formato europeo: 1.234,56
    if "," in text and "." in text:
        if text.rfind(",") > text.rfind("."):
            text = text.replace(".", "").replace(",", ".")
        else:
            text = text.replace(",", "")
    elif "," in text:
        text = text.replace(",", ".")
    try:
        return max(0.0, float(text))
    except Exception:
        return 0.0


def format_amount_value(value: Any) -> str:
    amount = parse_amount_value(value)
    if amount <= 0:
        return ""
    return f"{amount:.2f}"


def deduction_type_label(value: Any, language: str = "es") -> str:
    key = safe_str(value) or "none"
    t = TEXT[language]
    return {
        "none": t["no_deduction"],
        "partial": t["partial_deduction"],
        "total": t["full_deduction"],
    }.get(key, key)


def claim_deduction_amount(claim: Dict[str, Any]) -> float:
    return parse_amount_value(claim.get("deduction_amount", 0))


def claim_deduction_reason(claim: Dict[str, Any], language: str = "es") -> str:
    return comment_for_language(claim.get("deduction_reason", ""), language)


def has_deduction(claim: Dict[str, Any]) -> bool:
    return safe_str(claim.get("deduction_type", "none")) != "none" or claim_deduction_amount(claim) > 0


def mark_claim_as_not_warranty(claim: Dict[str, Any]) -> None:
    """Marca la claim como no garantía: todos los apartados puntuables a 0 y deducción total."""
    for check in ALL_SCORING_CHECKS:
        evaluation = claim.setdefault("evaluations", {}).setdefault(check.key, new_evaluation(check))
        evaluation["option_key"] = "nok"
        evaluation["status"] = "NOK"
        evaluation["points"] = 0
        evaluation["max_points"] = check.max_points
        if not safe_str(evaluation.get("comment", "")):
            evaluation["comment"] = "No corresponde a una reparación cubierta por garantía."

    claim["deduction_type"] = "total"
    amount = parse_amount_value(claim.get("amount", ""))
    claim["deduction_amount"] = amount
    claim["deduction_reason"] = "No corresponde a una reparación cubierta por garantía. Deducción del importe completo de la claim."
    claim["general_comment"] = build_general_comment_from_observations(claim, include_campaigns=False, language="es") or claim["deduction_reason"]


def mark_claim_as_not_warranty_callback(claim_key: str) -> None:
    claim = st.session_state.claims.get(claim_key)
    if not claim:
        return
    mark_claim_as_not_warranty(claim)
    # Sincronizar widgets de selectbox/text_area antes del siguiente rerun.
    for check in ALL_SCORING_CHECKS:
        st.session_state[f"select_{claim_key}_{check.key}"] = option_display(option_from_key(check, "nok"), "es")
        st.session_state[f"comment_{claim_key}_{check.key}"] = claim["evaluations"][check.key].get("comment", "")
    st.session_state[f"deduction_type_{claim_key}"] = "total"
    st.session_state[f"deduction_amount_{claim_key}"] = parse_amount_value(claim.get("amount", ""))
    st.session_state[f"deduction_reason_{claim_key}"] = claim.get("deduction_reason", "")
    st.session_state[f"general_comment_{claim_key}"] = claim.get("general_comment", "")

def get_dealer_options(current_value: str = "") -> List[str]:
    current_value = safe_str(current_value)
    options = [""] + ACTIVE_DEALERS.copy()
    if current_value and current_value not in options:
        options.insert(1, current_value)
    return options


def format_dealer_option(value: str) -> str:
    return "Selecciona dealer" if not safe_str(value) else value


def option_labels(check: AuditCheck, language: str = "es") -> List[str]:
    return [option.label_es if language == "es" else option.label_en for option in check.options]


def option_display(option: AuditOption, language: str = "es") -> str:
    return option.label_es if language == "es" else option.label_en


def check_label(check: AuditCheck, language: str = "es") -> str:
    return check.label_es if language == "es" else check.label_en


def check_guidance(check: AuditCheck, language: str = "es") -> str:
    return check.guidance_es if language == "es" else check.guidance_en


def block_label(block_key: str, language: str = "es") -> str:
    if block_key == "document":
        return TEXT[language]["document_section"]
    if block_key == "old_parts":
        return TEXT[language]["old_parts_section"]
    return safe_str(block_key)


def option_from_label(check: AuditCheck, label: str) -> AuditOption:
    for option in check.options:
        if label in (option.label_es, option.label_en):
            return option
    return PENDING


def option_from_key(check: AuditCheck, key: str) -> AuditOption:
    for option in check.options:
        if option.key == key:
            return option
    return PENDING


def option_from_points(check: AuditCheck, raw_points: Any) -> AuditOption:
    if raw_points is None or safe_str(raw_points) == "":
        return PENDING
    try:
        numeric = int(float(str(raw_points).replace(",", ".")))
    except Exception:
        return PENDING
    for option in check.options:
        if option.points == numeric:
            return option
    return PENDING


def infer_claim_ids(value: Any) -> Tuple[str, str]:
    value = safe_str(value)
    if not value:
        return "", ""
    normalized = value.upper().replace(" ", "")
    if normalized.startswith("CO"):
        return value, ""
    return "", value


def claim_identifier(claim: Dict[str, Any], language: str = "es") -> str:
    if language == "es":
        return safe_str(claim.get("local_claim_no")) or safe_str(claim.get("claim_no")) or safe_str(claim.get("hq_claim_no"))
    return safe_str(claim.get("hq_claim_no")) or safe_str(claim.get("claim_no")) or safe_str(claim.get("local_claim_no"))


def claim_other_identifier(claim: Dict[str, Any], language: str = "es") -> str:
    if language == "es":
        return safe_str(claim.get("hq_claim_no"))
    return safe_str(claim.get("local_claim_no"))


def make_claim_key(local_claim_no: str, hq_claim_no: str, fallback: str = "") -> str:
    local = safe_str(local_claim_no)
    hq = safe_str(hq_claim_no)
    fallback = safe_str(fallback)
    if local and hq:
        return f"{local}__{hq}"
    return local or hq or fallback


def result_label(points: Any, language: str = "es") -> str:
    t = TEXT[language]
    try:
        p = float(points)
    except Exception:
        return "Pendiente" if language == "es" else "Pending"
    if p >= 90:
        return t["excellent"]
    if p >= 80:
        return t["correct"]
    if p >= 60:
        return t["improvable"]
    return t["critical"]


# =============================================================================
# TRADUCCIÓN LOCAL DE OBSERVACIONES
# =============================================================================

# Glosario sencillo para que los exports EN no arrastren comentarios en español.
# No usa IA ni servicios externos: es determinista, rápido y seguro para Streamlit Cloud.
_COMMENT_TRANSLATION_PATTERNS: List[Tuple[str, str]] = [
    # Etiquetas y frases completas antes de sustituir palabras sueltas.
    (r"Pedido piezas / albarán", "Parts order / delivery note"),
    (r"Pedido piezas / albaran", "Parts order / delivery note"),
    (r"Pieza causa correcta", "Correct causal part"),
    (r"Fecha/hora envío Claim", "Claim submission and repair date"),
    (r"Fecha/hora envio Claim", "Claim submission and repair date"),
    (r"Gestión piezas viejas", "Old parts management"),
    (r"Gestion piezas viejas", "Old parts management"),
    (r"Etiquetado pieza vieja", "Old part labelling"),
    (r"Info tipo fallo pieza causa", "Causal part failure information"),
    (r"Destrucción pieza vieja", "Old part destruction"),
    (r"Destruccion pieza vieja", "Old part destruction"),
    (r"Certificado destrucción piezas viejas", "Old parts destruction certificate"),
    (r"Certificado destruccion piezas viejas", "Old parts destruction certificate"),
    (r"\bpedido piezas / albarán\b", "parts order / delivery note"),
    (r"\bpedido piezas / albaran\b", "parts order / delivery note"),
    (r"\bpieza causa correcta\b", "correct causal part"),
    (r"\bfecha/hora envío claim\b", "claim submission and repair date"),
    (r"\bfecha/hora envio claim\b", "claim submission and repair date"),
    (r"\bgestión piezas viejas\b", "old parts management"),
    (r"\bgestion piezas viejas\b", "old parts management"),
    (r"\betiquetado pieza vieja\b", "old part labelling"),
    (r"\binfo tipo fallo pieza causa\b", "causal part failure information"),
    (r"\bdestrucción pieza vieja\b", "old part destruction"),
    (r"\bdestruccion pieza vieja\b", "old part destruction"),
    (r"\bcertificado destrucción piezas viejas\b", "old parts destruction certificate"),
    (r"\bcertificado destruccion piezas viejas\b", "old parts destruction certificate"),
    (r"\bmano de obra\b", "labour"),
    (r"\bmaterial auxiliar\b", "auxiliary material"),
    (r"\bevidencias\b", "evidence"),
    (r"\bpieza causa\b", "causal part"),
    # Frases completas antes de sustituir palabras sueltas.
    (r"\bno corresponde a una reparación cubierta por garantía\b", "this does not correspond to a repair covered by warranty"),
    (r"\bno corresponde a una reparacion cubierta por garantia\b", "this does not correspond to a repair covered by warranty"),
    (r"\bdeducción del importe completo de la claim\b", "full deduction of the claim amount"),
    (r"\bdeduccion del importe completo de la claim\b", "full deduction of the claim amount"),
    (r"\bno se adjuntan\b", "are not attached"),
    (r"\bno se adjunta\b", "is not attached"),
    (r"\bno adjuntan\b", "do not attach"),
    (r"\bno adjunta\b", "does not attach"),
    (r"\bfaltan\b", "missing"),
    (r"\bfalta\b", "missing"),
    (r"\bsin\b", "without"),
    (r"\bno hay\b", "there is no"),
    (r"\bno existe\b", "does not exist"),
    (r"\bno procede\b", "not applicable"),
    (r"\bno aplica\b", "not applicable"),
    (r"\bincompleto\b", "incomplete"),
    (r"\bincompleta\b", "incomplete"),
    (r"\bparcial\b", "partial"),
    (r"\btotal\b", "full"),
    (r"\bincorrecto\b", "incorrect"),
    (r"\bincorrecta\b", "incorrect"),
    (r"\bcorrecto\b", "correct"),
    (r"\bcorrecta\b", "correct"),
    (r"\binsuficiente\b", "insufficient"),
    (r"\binsuficientes\b", "insufficient"),
    (r"\bno conforme\b", "non-compliant"),
    (r"\bconforme\b", "compliant"),
    (r"\bno queda justificado\b", "is not justified"),
    (r"\bno justificado\b", "not justified"),
    (r"\bno justificada\b", "not justified"),
    (r"\bjustificar\b", "justify"),
    (r"\bjustificado\b", "justified"),
    (r"\bjustificada\b", "justified"),
    (r"\bfuera de plazo\b", "outside the allowed deadline"),
    (r"\bplazo\b", "deadline"),
    (r"\bplazos\b", "deadlines"),
    (r"\bfecha de envio\b", "submission date"),
    (r"\bfecha de envío\b", "submission date"),
    (r"\bfecha de reparacion\b", "repair date"),
    (r"\bfecha de reparación\b", "repair date"),
    (r"\borden de reparacion\b", "repair order"),
    (r"\borden de reparación\b", "repair order"),
    (r"\bla or\b", "the repair order"),
    (r"\bel or\b", "the repair order"),
    (r"\bor previa\b", "previous repair order"),
    (r"\bor\b", "repair order"),
    (r"\breparacion\b", "repair"),
    (r"\breparación\b", "repair"),
    (r"\breparaciones\b", "repairs"),
    (r"\bpedido de piezas\b", "parts order"),
    (r"\balbaran\b", "delivery note"),
    (r"\balbarán\b", "delivery note"),
    (r"\bevidencias\b", "evidence"),
    (r"\bevidencia\b", "evidence"),
    (r"\bfotos\b", "photos"),
    (r"\bfoto\b", "photo"),
    (r"\bpieza causa\b", "causal part"),
    (r"\bpiezas viejas\b", "old parts"),
    (r"\bpieza vieja\b", "old part"),
    (r"\bpieza sustituida\b", "replaced part"),
    (r"\bpieza nueva\b", "new part"),
    (r"\bpiezas\b", "parts"),
    (r"\bpieza\b", "part"),
    (r"\bmano de obra adicional\b", "additional labour"),
    (r"\bmano de obra\b", "labour"),
    (r"\bmaterial auxiliar\b", "auxiliary material"),
    (r"\bcoste auxiliar\b", "auxiliary cost"),
    (r"\bcostes adicionales\b", "additional costs"),
    (r"\bcoste\b", "cost"),
    (r"\bimporte\b", "amount"),
    (r"\bdeduccion\b", "deduction"),
    (r"\bdeducción\b", "deduction"),
    (r"\bdeducir\b", "deduct"),
    (r"\bno es garantia\b", "not covered by warranty"),
    (r"\bno es garantía\b", "not covered by warranty"),
    (r"\bno corresponde a una reparacion cubierta por garantia\b", "does not correspond to a repair covered by warranty"),
    (r"\bno corresponde a una reparación cubierta por garantía\b", "does not correspond to a repair covered by warranty"),
    (r"\btiempo adicional\b", "additional time"),
    (r"\btiempo extra\b", "extra time"),
    (r"\bbaremo\b", "standard time"),
    (r"\boperacion\b", "operation"),
    (r"\boperación\b", "operation"),
    (r"\boperaciones\b", "operations"),
    (r"\bdiagnostico\b", "diagnosis"),
    (r"\bdiagnóstico\b", "diagnosis"),
    (r"\baveria\b", "failure"),
    (r"\bavería\b", "failure"),
    (r"\bsintoma\b", "symptom"),
    (r"\bsíntoma\b", "symptom"),
    (r"\bsolucion\b", "solution"),
    (r"\bsolución\b", "solution"),
    (r"\bfirma del cliente\b", "customer signature"),
    (r"\bfirma\b", "signature"),
    (r"\bcliente\b", "customer"),
    (r"\bvehiculo\b", "vehicle"),
    (r"\bvehículo\b", "vehicle"),
    (r"\bkilometraje\b", "mileage"),
    (r"\betiqueta\b", "label"),
    (r"\betiquetado\b", "labelling"),
    (r"\bdestruccion\b", "destruction"),
    (r"\bdestrucción\b", "destruction"),
    (r"\bcertificado\b", "certificate"),
    (r"\binforme\b", "report"),
    (r"\bcampañas\b", "campaigns"),
    (r"\bcampana\b", "campaign"),
    (r"\bcampaña\b", "campaign"),
    (r"\bpendiente\b", "pending"),
    (r"\bpendientes\b", "pending"),
    (r"\brevisar\b", "review"),
    (r"\badjuntar\b", "attach"),
    (r"\badjunta\b", "attaches"),
    (r"\badjuntado\b", "attached"),
    (r"\badjuntada\b", "attached"),
    (r"\bsubir\b", "upload"),
    (r"\barchivar\b", "archive"),
    (r"\bcomprobar\b", "check"),
    (r"\bverificar\b", "verify"),
    (r"\bcoherente\b", "consistent"),
    (r"\bcoherencia\b", "consistency"),
    (r"\bcausa\b", "cause"),
    (r"\bfallo\b", "failure"),
    (r"\btipo de fallo\b", "failure type"),
    (r"\bdatos\b", "data"),
    (r"\breclamacion\b", "claim"),
    (r"\breclamación\b", "claim"),
    (r"\bgarantia\b", "warranty"),
    (r"\bgarantía\b", "warranty"),
    (r"\bde la\b", "of the"),
    (r"\bde las\b", "of the"),
    (r"\bde los\b", "of the"),
    (r"\bdel\b", "of the"),
    (r"\bde\b", "of"),
]


def looks_like_english(text: str) -> bool:
    sample = normalize_text(text)
    if not sample:
        return True
    spanish_markers = [
        " falta ", " faltan ", " evidencia ", " evidencias ", " pieza ", " piezas ",
        " reparacion ", " garantia ", " reclamacion ", " albaran ", " fecha ",
        " no se ", " mano de obra ", " material auxiliar ", " certificado ",
    ]
    english_markers = [
        " missing ", " evidence ", " repair ", " warranty ", " claim ", " labour ",
        " part ", " parts ", " delivery note ", " certificate ", " deadline ",
    ]
    padded = f" {sample} "
    spanish_hits = sum(marker in padded for marker in spanish_markers)
    english_hits = sum(marker in padded for marker in english_markers)
    return english_hits > 0 and spanish_hits == 0


def translate_comment_to_english(text: Any) -> str:
    """
    Traduce observaciones habituales de auditoría al inglés usando un glosario local.
    Si el texto ya parece inglés, lo deja intacto. No llama a APIs externas.
    """
    original = safe_str(text)
    if not original or looks_like_english(original):
        return original

    translated_lines = []
    for line in original.splitlines():
        line = safe_str(line)
        if not line:
            translated_lines.append("")
            continue

        translated = line
        # Primero, expresiones completas muy comunes.
        exact_map = {
            "ok": "OK",
            "no aplica": "Not applicable",
            "pendiente": "Pending",
            "sin observación específica.": "No specific observation.",
            "sin observacion especifica.": "No specific observation.",
            "falta evidencia de diagnostico": "Missing diagnosis evidence.",
            "falta evidencia de diagnostico.": "Missing diagnosis evidence.",
            "faltan evidencias de diagnostico": "Missing diagnosis evidence.",
            "faltan evidencias de diagnostico.": "Missing diagnosis evidence.",
            "tiempo adicional no justificado": "Additional time is not justified.",
            "tiempo adicional no justificado.": "Additional time is not justified.",
            "mano de obra adicional no justificada": "Additional labour is not justified.",
            "mano de obra adicional no justificada.": "Additional labour is not justified.",
            "falta firma del cliente": "Missing customer signature.",
            "falta firma del cliente.": "Missing customer signature.",
            "no se adjunta la or": "The repair order is not attached.",
            "no se adjunta la or.": "The repair order is not attached.",
            "no es garantia": "Not covered by warranty.",
            "no es garantía": "Not covered by warranty.",
            "deduccion parcial": "Partial deduction.",
            "deducción parcial": "Partial deduction.",
            "deduccion total": "Full deduction.",
            "deducción total": "Full deduction.",
            "no corresponde a una reparacion cubierta por garantia": "This does not correspond to a repair covered by warranty.",
            "no corresponde a una reparación cubierta por garantía": "This does not correspond to a repair covered by warranty.",
        }
        normalized_line = normalize_text(line)
        if normalized_line in exact_map:
            translated_lines.append(exact_map[normalized_line])
            continue

        # Sustituciones por glosario. Se hacen con ignorecase y manteniendo lo demás.
        for pattern, replacement in _COMMENT_TRANSLATION_PATTERNS:
            translated = re.sub(pattern, replacement, translated, flags=re.IGNORECASE)

        # Limpieza mínima de dobles espacios y mayúscula inicial.
        translated = re.sub(r"\s+", " ", translated).strip()
        if translated and translated[0].islower():
            translated = translated[0].upper() + translated[1:]
        translated_lines.append(translated)

    return "\n".join(translated_lines).strip()


def comment_for_language(text: Any, language: str = "es") -> str:
    comment = safe_str(text)
    if language == "en":
        return translate_comment_to_english(comment)
    return comment


def status_for_language(status: Any, language: str = "es") -> str:
    value = safe_str(status)
    if language != "en":
        return value
    normalized = normalize_text(value)
    return {
        "pendiente": "Pending",
        "ok": "OK",
        "parcial": "Partial",
        "nok": "NOK",
        "n a": "N/A",
        "na": "N/A",
    }.get(normalized, value)


def pending_items_for_language(claim: Dict[str, Any], language: str = "es") -> List[str]:
    pending = []
    for check in ALL_SCORING_CHECKS:
        evaluation = claim.get("evaluations", {}).get(check.key, {})
        if evaluation.get("points") is None:
            pending.append(check_label(check, language))
    return pending


# =============================================================================
# MODELO DE DATOS
# =============================================================================


def new_evaluation(check: AuditCheck, prefill_points: Any = None) -> Dict[str, Any]:
    """
    Crea la evaluación inicial de un apartado.

    Para agilizar la auditoría, por defecto se considera OK / puntuación máxima.
    El auditor solo tiene que tocar los apartados con desviación.

    Si llega una puntuación existente desde una plantilla/JSON, se respeta.
    """
    if prefill_points is None or safe_str(prefill_points) == "":
        option = option_from_key(check, "ok")
    else:
        option = option_from_points(check, prefill_points)

    return {
        "option_key": option.key,
        "status": option.status,
        "points": option.points,
        "max_points": check.max_points,
        "comment": "",
    }


def empty_claim_record(claim_no: str = "", local_claim_no: str = "", hq_claim_no: str = "") -> Dict[str, Any]:
    inferred_local, inferred_hq = infer_claim_ids(claim_no)
    local = safe_str(local_claim_no) or inferred_local
    hq = safe_str(hq_claim_no) or inferred_hq
    internal_claim_no = safe_str(claim_no) or local or hq
    return {
        "claim_no": internal_claim_no,
        "local_claim_no": local,
        "hq_claim_no": hq,
        "dealer": "",
        "vin": "",
        "model": "",
        "amount": "",
        "deduction_type": "none",
        "deduction_amount": 0.0,
        "deduction_reason": "",
        "repair_date": "",
        "submission_date": "",
        "general_comment": "",
        "internal_comment": "",
        "evaluations": {check.key: new_evaluation(check) for check in ALL_CHECKS},
    }


def calculate_claim_score(claim: Dict[str, Any]) -> Dict[str, Any]:
    doc_points = 0
    old_points = 0
    pending = []
    lost_by_area = []

    for check in DOCUMENT_CHECKS:
        evaluation = claim["evaluations"].get(check.key, {})
        points = evaluation.get("points")
        if points is None:
            pending.append(check.label_es)
            continue
        doc_points += int(points)
        lost = check.max_points - int(points)
        if lost > 0:
            lost_by_area.append((check, lost, evaluation.get("status", "")))

    for check in OLD_PARTS_CHECKS:
        evaluation = claim["evaluations"].get(check.key, {})
        points = evaluation.get("points")
        if points is None:
            pending.append(check.label_es)
            continue
        old_points += int(points)
        lost = check.max_points - int(points)
        if lost > 0:
            lost_by_area.append((check, lost, evaluation.get("status", "")))

    total_points = doc_points + old_points
    return {
        "doc_points": doc_points,
        "old_points": old_points,
        "total_points": total_points,
        "success_percent": total_points,
        "pending": pending,
        "completed": len(pending) == 0,
        "lost_by_area": lost_by_area,
    }


def calculate_audit_score(claims: Dict[str, Dict[str, Any]]) -> Dict[str, Any]:
    if not claims:
        return {"claims": 0, "completed_claims": 0, "doc_points": 0, "old_points": 0, "total_points": 0, "max_points": 0, "success_percent": 0}
    doc_points = old_points = total_points = completed = 0
    for claim in claims.values():
        score = calculate_claim_score(claim)
        doc_points += score["doc_points"]
        old_points += score["old_points"]
        total_points += score["total_points"]
        completed += int(bool(score["completed"]))
    max_points = len(claims) * MAX_TOTAL_POINTS
    return {
        "claims": len(claims),
        "completed_claims": completed,
        "doc_points": doc_points,
        "old_points": old_points,
        "total_points": total_points,
        "max_points": max_points,
        "success_percent": (total_points / max_points * 100) if max_points else 0,
    }


def normalize_loaded_evaluation(check: AuditCheck, raw: Any) -> Dict[str, Any]:
    base = new_evaluation(check)
    if not isinstance(raw, dict):
        return base
    option_key = safe_str(raw.get("option_key", ""))
    status = safe_str(raw.get("status", ""))
    points = raw.get("points", None)
    if option_key:
        option = option_from_key(check, option_key)
    elif points is not None:
        option = option_from_points(check, points)
    else:
        # Si no hay evaluación guardada, tratamos el apartado como OK por defecto.
        option = option_from_key(check, "ok")
    # Mantener pendiente si venía explícitamente pendiente.
    final_points = None if status.lower().startswith("pend") and points is None else option.points
    return {
        "option_key": option.key,
        "status": status or option.status,
        "points": final_points,
        "max_points": check.max_points,
        "comment": safe_str(raw.get("comment", "")),
    }


def normalize_loaded_claim(raw_claim: Any, fallback: str = "") -> Optional[Dict[str, Any]]:
    if not isinstance(raw_claim, dict):
        return None
    claim_no = safe_str(raw_claim.get("claim_no", fallback))
    local = safe_str(raw_claim.get("local_claim_no", ""))
    hq = safe_str(raw_claim.get("hq_claim_no", ""))
    if not any([claim_no, local, hq]):
        return None
    claim = empty_claim_record(claim_no, local, hq)
    for field in ["dealer", "vin", "model", "amount", "deduction_type", "deduction_amount", "deduction_reason", "repair_date", "submission_date", "general_comment", "internal_comment"]:
        claim[field] = safe_str(raw_claim.get(field, claim.get(field, "")))
    raw_evaluations = raw_claim.get("evaluations", {})
    if isinstance(raw_evaluations, dict):
        for check in ALL_CHECKS:
            if check.key in raw_evaluations:
                claim["evaluations"][check.key] = normalize_loaded_evaluation(check, raw_evaluations[check.key])
    return claim


def serialize_audit_workfile(claims: Dict[str, Dict[str, Any]], audit_name: str, dealer: str, auditor: str, audit_date_value: Any = None) -> bytes:
    sync_all_claims_from_widget_state(claims)
    audit_date = parse_audit_date(audit_date_value).isoformat()
    payload = {
        "file_type": "warranty_audit_workfile",
        "version": 3,
        "saved_at": datetime.now().isoformat(timespec="seconds"),
        "audit": {
            "audit_name": audit_name or "",
            "dealer": dealer or "",
            "auditor": auditor or "",
            "audit_date": audit_date,
        },
        "claims": claims,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def load_audit_workfile(uploaded_file) -> Tuple[Dict[str, Dict[str, Any]], str, str, str, date]:
    content = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
    if isinstance(content, bytes):
        content = content.decode("utf-8")
    payload = json.loads(content)
    if not isinstance(payload, dict) or payload.get("file_type") != "warranty_audit_workfile":
        raise ValueError("El archivo no parece una auditoría de trabajo generada por esta app.")
    audit = payload.get("audit", {}) if isinstance(payload.get("audit", {}), dict) else {}
    raw_claims = payload.get("claims", {})
    claims: Dict[str, Dict[str, Any]] = {}
    iterable = raw_claims.items() if isinstance(raw_claims, dict) else []
    for fallback_key, raw_claim in iterable:
        claim = normalize_loaded_claim(raw_claim, fallback_key)
        if claim:
            claims[make_claim_key(claim.get("local_claim_no"), claim.get("hq_claim_no"), claim.get("claim_no"))] = claim
    if not claims:
        raise ValueError("El archivo se pudo abrir, pero no contiene claims válidas.")
    return (
        claims,
        safe_str(audit.get("audit_name", "Auditoría garantías")) or "Auditoría garantías",
        safe_str(audit.get("dealer", "")),
        safe_str(audit.get("auditor", "")),
        parse_audit_date(audit.get("audit_date", datetime.now().date().isoformat())),
    )


# =============================================================================
# IMPORTACIÓN DE EXCEL
# =============================================================================


def find_column(normalized_cols: Dict[str, Any], candidates: List[str]) -> Optional[Any]:
    normalized_candidates = [normalize_text(item) for item in candidates]
    for candidate in normalized_candidates:
        if candidate in normalized_cols:
            return normalized_cols[candidate]
    # segunda pasada por inclusión flexible
    for candidate in normalized_candidates:
        for col_norm, col_real in normalized_cols.items():
            if candidate and candidate in col_norm:
                return col_real
    return None


def upsert_claim(claims: Dict[str, Dict[str, Any]], claim_no: str = "", local_claim_no: str = "", hq_claim_no: str = "") -> Dict[str, Any]:
    local = safe_str(local_claim_no)
    hq = safe_str(hq_claim_no)
    if not local and not hq and claim_no:
        local, hq = infer_claim_ids(claim_no)
    key = make_claim_key(local, hq, claim_no)
    if not key:
        raise ValueError("Claim vacía")
    claim = claims.setdefault(key, empty_claim_record(claim_no or local or hq, local, hq))
    if local:
        claim["local_claim_no"] = local
    if hq:
        claim["hq_claim_no"] = hq
    if not safe_str(claim.get("claim_no", "")):
        claim["claim_no"] = local or hq or claim_no
    return claim


def read_claims_from_uploaded_excel(uploaded_file) -> Dict[str, Dict[str, Any]]:
    xls = pd.ExcelFile(uploaded_file)
    claims: Dict[str, Dict[str, Any]] = {}

    if "Claim document checklist I" in xls.sheet_names:
        doc_df = pd.read_excel(uploaded_file, sheet_name="Claim document checklist I", header=None)
        doc_columns = {
            "doc_or": 2,
            "doc_parts_order": 3,
            "doc_previous_or": 4,
            "doc_evidence": 5,
            "doc_causal_part": 6,
            "doc_labor": 7,
            "doc_aux_material": 8,
            "doc_dates": 9,
            "doc_vin": 10,
        }
        for row_idx in range(3, len(doc_df)):
            raw_claim = safe_str(doc_df.iloc[row_idx, 1] if doc_df.shape[1] > 1 else "")
            if not raw_claim or raw_claim.lower() == "nan":
                continue
            local, hq = infer_claim_ids(raw_claim)
            claim = upsert_claim(claims, raw_claim, local, hq)
            for key, col_idx in doc_columns.items():
                if col_idx < doc_df.shape[1]:
                    claim["evaluations"][key] = new_evaluation(CHECK_BY_KEY[key], doc_df.iloc[row_idx, col_idx])
            if doc_df.shape[1] > 15:
                comment = safe_str(doc_df.iloc[row_idx, 15])
                if comment:
                    claim["general_comment"] = comment

    if "Claim old parts checklist II" in xls.sheet_names:
        old_df = pd.read_excel(uploaded_file, sheet_name="Claim old parts checklist II", header=None)
        old_columns = {
            "old_management": 2,
            "old_label": 3,
            "old_causal_part": 4,
            "old_failure_info": 5,
            "old_destruction": 6,
            "old_destruction_certificate": 7,
        }
        for row_idx in range(3, len(old_df)):
            raw_claim = safe_str(old_df.iloc[row_idx, 1] if old_df.shape[1] > 1 else "")
            if not raw_claim or raw_claim.lower() == "nan":
                continue
            local, hq = infer_claim_ids(raw_claim)
            claim = upsert_claim(claims, raw_claim, local, hq)
            for key, col_idx in old_columns.items():
                if col_idx < old_df.shape[1]:
                    claim["evaluations"][key] = new_evaluation(CHECK_BY_KEY[key], old_df.iloc[row_idx, col_idx])
            if old_df.shape[1] > 10:
                old_comment = safe_str(old_df.iloc[row_idx, 10])
                if old_comment:
                    existing = safe_str(claim.get("general_comment", ""))
                    claim["general_comment"] = (existing + "\n" + old_comment).strip() if existing else old_comment

    if not claims:
        df = pd.read_excel(uploaded_file, sheet_name=xls.sheet_names[0])
        normalized_cols = {normalize_text(col): col for col in df.columns}
        local_col = find_column(normalized_cols, [
            "Claim España", "Claim ES", "Spanish claim", "Dealer claim", "Local claim", "Claim No.", "Claim No", "Garantía", "Garantia", "CO", "Claim",
        ])
        hq_col = find_column(normalized_cols, [
            "HQ Claim", "HQ Claim No", "IDMS Claim", "IDMs Claim", "TAC", "TAC No", "Identificador HQ", "Claim HQ", "2810",
        ])
        generic_col = find_column(normalized_cols, ["Claim No.", "Claim No", "Claim", "claim_number", "claim number", "Garantía", "Garantia"])
        if local_col is None and hq_col is None and generic_col is None:
            raise ValueError("No encuentro columnas de claim. Usa al menos Claim España/CO o HQ Claim/IDMS.")

        meta_cols = {
            "dealer": find_column(normalized_cols, ["dealer", "concesionario", "service dealer"]),
            "vin": find_column(normalized_cols, ["vin", "chassis", "bastidor"]),
            "model": find_column(normalized_cols, ["model", "modelo"]),
            "amount": find_column(normalized_cols, ["amount", "importe", "coste", "total"]),
            "repair_date": find_column(normalized_cols, ["repair date", "fecha reparacion", "fecha reparación"]),
            "submission_date": find_column(normalized_cols, ["submission date", "fecha envio", "fecha envío"]),
        }

        for _, row in df.iterrows():
            local_value = safe_str(row.get(local_col, "")) if local_col is not None else ""
            hq_value = safe_str(row.get(hq_col, "")) if hq_col is not None else ""
            generic_value = safe_str(row.get(generic_col, "")) if generic_col is not None else ""

            if not local_value and not hq_value and generic_value:
                local_value, hq_value = infer_claim_ids(generic_value)
            if not local_value and not hq_value:
                continue

            claim = upsert_claim(claims, generic_value or local_value or hq_value, local_value, hq_value)
            for field, source_col in meta_cols.items():
                if source_col is not None:
                    claim[field] = safe_str(row.get(source_col, ""))

    return claims


# =============================================================================
# COMENTARIOS Y PLAN DE ACCIÓN
# =============================================================================


def evaluation_comment(claim: Dict[str, Any], check: AuditCheck) -> str:
    return safe_str(claim.get("evaluations", {}).get(check.key, {}).get("comment", ""))


def build_general_comment_from_observations(claim: Dict[str, Any], include_campaigns: bool = True, language: str = "es") -> str:
    # Campañas eliminadas del flujo: el comentario general se genera solo desde apartados auditables.
    checks = ALL_SCORING_CHECKS
    lines = []
    for check in checks:
        comment = comment_for_language(evaluation_comment(claim, check), language)
        if not comment:
            continue
        label = check_label(check, language)
        if comment.lower().lstrip().startswith(label.lower() + ":"):
            lines.append(comment)
        else:
            lines.append(f"{label}: {comment}")
    return "\n".join(lines).strip()


def auto_general_comment(claim: Dict[str, Any], language: str = "es") -> str:
    """
    Comentario general calculado automáticamente a partir de las observaciones
    de los apartados. Si no hay observaciones, usa el motivo de deducción si existe.
    """
    generated = build_general_comment_from_observations(
        claim,
        include_campaigns=False,
        language=language,
    )
    if generated:
        return generated
    return claim_deduction_reason(claim, language)


def refresh_claim_general_comment(claim: Dict[str, Any]) -> str:
    """
    Guarda en la claim el comentario general en español calculado automáticamente.
    Se conserva en el JSON para que el progreso sea transparente, pero no se edita a mano.
    """
    comment = auto_general_comment(claim, "es")
    claim["general_comment"] = comment
    return comment


def get_streamlit_state_value(key: str, default: Any = "") -> Any:
    """Lee session_state de forma segura incluso si la función se llama fuera de Streamlit."""
    try:
        return st.session_state.get(key, default)
    except Exception:
        return default


def get_ai_output(key: str) -> str:
    value = get_streamlit_state_value(key, "")
    return safe_str(value)


def get_ai_claim_comment(claim: Dict[str, Any], language: str = "es") -> str:
    """Devuelve comentario de claim pulido/traducido por IA si existe."""
    if language == "en":
        try:
            comments = st.session_state.get("ai_claim_comments_en", {}) or {}
        except Exception:
            comments = {}
        identifiers = [
            claim_identifier(claim, "en"),
            safe_str(claim.get("hq_claim_no", "")),
            safe_str(claim.get("local_claim_no", "")),
            safe_str(claim.get("claim_no", "")),
        ]
        for identifier in identifiers:
            if identifier and identifier in comments and safe_str(comments[identifier]):
                return safe_str(comments[identifier])
    return ""


def general_comment_for_export(claim: Dict[str, Any], language: str = "es") -> str:
    """Comentario usado en exports. En inglés aprovecha traducciones IA si ya se han generado."""
    ai_comment = get_ai_claim_comment(claim, language)
    if ai_comment:
        return ai_comment
    return auto_general_comment(claim, language)


def get_ai_report_for_language(language: str = "es") -> str:
    return get_ai_output("ai_report_es" if language == "es" else "ai_report_en")


def get_ai_action_plan_for_language(language: str = "es") -> str:
    return get_ai_output("ai_action_plan_es" if language == "es" else "ai_action_plan_en")


def get_secret_value(*names: str) -> str:
    """Busca una clave en Streamlit Secrets o variables de entorno."""
    for name in names:
        try:
            value = st.secrets.get(name, "")
            if value:
                return safe_str(value)
        except Exception:
            pass
        value = os.environ.get(name, "")
        if value:
            return safe_str(value)
    return ""


def build_ai_audit_payload(claims: Dict[str, Dict[str, Any]], audit_name: str, dealer: str, auditor: str) -> Dict[str, Any]:
    """Prepara un contexto compacto para Gemini. La IA solo redacta/resume; no recalcula puntos."""
    sync_all_claims_from_widget_state(claims)
    audit_score = calculate_audit_score(claims)
    action_rows_es = build_action_plan_rows(claims, "es")

    claim_payload = []
    for claim in claims.values():
        score = calculate_claim_score(claim)
        # calculate_claim_score incluye objetos AuditCheck dentro de lost_by_area,
        # y esos objetos no se pueden convertir directamente a JSON para Gemini.
        # Creamos una versión limpia y serializable del score para el prompt.
        score_for_ai = {
            "doc_points": int(score.get("doc_points", 0)),
            "old_points": int(score.get("old_points", 0)),
            "total_points": int(score.get("total_points", 0)),
            "success_percent": float(score.get("success_percent", 0) or 0),
            "pending": [safe_str(item) for item in score.get("pending", [])],
            "completed": bool(score.get("completed", False)),
            "lost_by_area": [
                {
                    "item_key": item_check.key,
                    "item_es": item_check.label_es,
                    "item_en": item_check.label_en,
                    "section": block_label(item_check.block_key, "es"),
                    "lost_points": int(lost_points),
                    "status": safe_str(status),
                }
                for item_check, lost_points, status in score.get("lost_by_area", [])
            ],
        }
        observations = []
        for check in ALL_SCORING_CHECKS:
            evaluation = claim.get("evaluations", {}).get(check.key, {})
            points = evaluation.get("points")
            lost = check.max_points if points is None else max(0, check.max_points - int(points or 0))
            comment = evaluation_comment(claim, check)
            if lost > 0 or comment:
                observations.append({
                    "section": block_label(check.block_key, "es"),
                    "item_key": check.key,
                    "item_es": check.label_es,
                    "item_en": check.label_en,
                    "status": evaluation.get("status", ""),
                    "points": points,
                    "max_points": check.max_points,
                    "lost_points": lost,
                    "observation_es": comment,
                })
        deduction = {
            "type": safe_str(claim.get("deduction_type", "none")) or "none",
            "amount": claim_deduction_amount(claim),
            "reason_es": safe_str(claim.get("deduction_reason", "")),
        }
        claim_payload.append({
            "claim_hq_id": safe_str(claim.get("hq_claim_no", "")),
            "claim_es_id": safe_str(claim.get("local_claim_no", "")),
            "dealer": safe_str(claim.get("dealer", "")) or dealer,
            "score": score_for_ai,
            "deduction": deduction,
            "automatic_general_comment_es": auto_general_comment(claim, "es"),
            "observations": observations,
        })

    return {
        "audit_name": audit_name,
        "dealer_general": dealer,
        "auditor": auditor,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "scoring_rule": {
            "document_checklist_max": MAX_DOCUMENT_POINTS,
            "old_parts_checklist_max": MAX_OLD_PARTS_POINTS,
            "total_max": MAX_TOTAL_POINTS,
            "not_applicable_rule": "N/A gives maximum points and does not penalize the score",
            "campaigns": "Not included",
        },
        "audit_score": audit_score,
        "main_improvement_rows_es": action_rows_es,
        "claims": claim_payload,
    }


def extract_json_object(text: str) -> Dict[str, Any]:
    """Intenta extraer JSON aunque el modelo lo devuelva dentro de ```json."""
    raw = safe_str(text)
    if not raw:
        return {}
    raw = re.sub(r"^```(?:json)?", "", raw.strip(), flags=re.IGNORECASE).strip()
    raw = re.sub(r"```$", "", raw.strip()).strip()
    try:
        return json.loads(raw)
    except Exception:
        pass
    start = raw.find("{")
    end = raw.rfind("}")
    if start >= 0 and end > start:
        try:
            return json.loads(raw[start:end + 1])
        except Exception:
            return {}
    return {}


def build_gemini_audit_prompt(payload: Dict[str, Any]) -> str:
    return f"""
You are an automotive warranty audit assistant for an NSC warranty department.

Your task is to create a professional executive audit report and action plan in Spanish and English.
Use ONLY the audit data provided below. Do not invent claims, scores, deductions or facts.
Do not recalculate the score. Use the provided scores exactly.

IMPORTANT STYLE REQUIREMENTS:
- Spanish text is for the dealer: clear, professional, direct and improvement-oriented.
- English text is for HQ: corporate, concise and formal.
- Summarize repeated observations into meaningful improvement areas.
- Do not output a raw list like "23 claims with deviation" as the main content.
- Explain what the pattern means and what action should be taken.
- Keep the tone constructive, not accusatory.
- Manual Spanish observations must be translated naturally into English in the English output.
- Scoring is rule-based; do not change it.
- Do not use excessive asterisks or bold Markdown markers. Keep the report readable in plain text too.

IMPORTANT CHECKLIST NAMING RULES:
- Respect the current checklist naming exactly.
- In Spanish, use the exact Spanish audited parameter names received in the audit data.
- In English, use the exact English audited parameter names received in the audit data.
- Do not rename audited parameters into generic alternatives.

ACTION PLAN STRUCTURE:
The action plan must follow the audit checklist structure, like an audit scorecard/action plan table.
Return action_plan_rows_es and action_plan_rows_en as arrays with ONE ROW FOR EACH audited parameter, in this exact order:
1. OR
2. Pedido piezas / albarán
3. OR previa
4. Evidencias
5. Pieza causa correcta
6. Mano de obra
7. Material auxiliar
8. Fecha/hora envío Claim
9. VIN
10. Gestión piezas viejas
11. Etiquetado pieza vieja
12. Pieza causa
13. Info tipo fallo pieza causa
14. Destrucción pieza vieja
15. Certificado destrucción piezas viejas

For English rows, use the exact English checklist names from the audit data.
Each row must contain:
- group: "Claim" for document checklist items, "Old part" for old parts checklist items.
- parameter: exact checklist parameter name.
- evaluation: "X" when there is any deviation, deduction impact or relevant observation for that parameter. Otherwise "V".
- observations: concise aggregated observation. Empty string if evaluation is "V".
- countermeasures: concise practical countermeasure. Empty string if evaluation is "V".

Return also action_plan_es and action_plan_en as readable Markdown tables with columns:
Audited parameters | Evaluation | Observations | Countermeasures

REPORT FORMAT:
- report_es and report_en must be highly readable.
- Use short plain headings and short bullet points.
- Recommended structure: Executive summary, Main findings, Main improvement areas, Conclusion.

Return ONLY valid JSON with this exact structure:
{{
  "report_es": "Texto completo del informe ejecutivo en español",
  "report_en": "Full executive report in English",
  "action_plan_es": "Tabla markdown del plan de acción en español",
  "action_plan_en": "Markdown table of the action plan in English",
  "action_plan_rows_es": [
    {{"group": "Claim", "parameter": "OR", "evaluation": "V", "observations": "", "countermeasures": ""}}
  ],
  "action_plan_rows_en": [
    {{"group": "Claim", "parameter": "Repair order", "evaluation": "V", "observations": "", "countermeasures": ""}}
  ],
  "claim_comments_en": {{
    "CLAIM_HQ_OR_ES_ID": "Professional English translation/summary of the claim observations"
  }}
}}

For claim_comments_en:
- Use the HQ claim ID when available.
- If the HQ ID is empty, use the Spanish claim ID.
- Include one entry for every claim that has observations or deductions.
- Keep each claim comment concise but clear.

Audit data JSON:
{json.dumps(payload, ensure_ascii=False, indent=2, default=json_default_serializer)}
""".strip()


def generate_gemini_audit_outputs(claims: Dict[str, Dict[str, Any]], audit_name: str, dealer: str, auditor: str, api_key: str, model: str = GEMINI_MODEL_DEFAULT) -> Tuple[bool, str]:
    """Llama a Gemini y guarda informe/plan ES-EN en session_state."""
    api_key = safe_str(api_key)
    if not api_key:
        return False, "No hay API key de Gemini. Añádela en Streamlit Secrets como GEMINI_API_KEY o pégala temporalmente en la app."

    try:
        from google import genai
        try:
            from google.genai import types
        except Exception:
            types = None
    except Exception as exc:
        return False, f"No está instalado el SDK de Gemini. Añade google-genai a requirements.txt. Detalle: {exc}"

    payload = build_ai_audit_payload(claims, audit_name, dealer, auditor)
    prompt = build_gemini_audit_prompt(payload)

    try:
        client = genai.Client(api_key=api_key)
        request_kwargs = {
            "model": model or GEMINI_MODEL_DEFAULT,
            "contents": prompt,
        }
        # Temperatura baja: más determinista, rápida y menos propensa a inventar formato.
        if types is not None:
            request_kwargs["config"] = types.GenerateContentConfig(temperature=0.15)
        else:
            request_kwargs["config"] = {"temperature": 0.15}
        response = client.models.generate_content(**request_kwargs)
        raw_text = getattr(response, "text", "") or ""
    except Exception as exc:
        return False, f"Gemini no ha podido generar el informe: {exc}"

    parsed = extract_json_object(raw_text)
    if not parsed:
        return False, "Gemini respondió, pero no se pudo interpretar como JSON. Prueba otra vez."

    st.session_state.ai_report_es = safe_str(parsed.get("report_es", ""))
    st.session_state.ai_report_en = safe_str(parsed.get("report_en", ""))
    st.session_state.ai_action_plan_es = safe_str(parsed.get("action_plan_es", ""))
    st.session_state.ai_action_plan_en = safe_str(parsed.get("action_plan_en", ""))

    rows_es = parsed.get("action_plan_rows_es", [])
    rows_en = parsed.get("action_plan_rows_en", [])
    st.session_state.ai_action_plan_rows_es = rows_es if isinstance(rows_es, list) else []
    st.session_state.ai_action_plan_rows_en = rows_en if isinstance(rows_en, list) else []

    claim_comments = parsed.get("claim_comments_en", {})
    st.session_state.ai_claim_comments_en = claim_comments if isinstance(claim_comments, dict) else {}
    st.session_state.ai_generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    return True, "Informe ejecutivo, plan de acción y traducciones EN generadas con Gemini."


def clear_ai_outputs() -> None:
    for key in [
        "ai_report_es", "ai_report_en", "ai_action_plan_es", "ai_action_plan_en",
        "ai_action_plan_rows_es", "ai_action_plan_rows_en",
        "ai_claim_comments_en", "ai_generated_at", "ai_model_used",
    ]:
        if key in st.session_state:
            del st.session_state[key]


def sync_claim_from_widget_state(claim: Dict[str, Any]) -> None:
    """
    Antes de calcular/exportar, recoge valores que ya estén en widgets Streamlit.

    Esto evita que un texto recién escrito en un text_area se quede fuera si el
    usuario pulsa descargar sin pasar antes por otra sección de la app.
    """
    try:
        session = st.session_state
    except Exception:
        refresh_claim_general_comment(claim)
        return

    claim_key = make_claim_key(
        claim.get("local_claim_no"),
        claim.get("hq_claim_no"),
        claim.get("claim_no"),
    )

    for field in ["hq_claim_no", "local_claim_no", "vin", "model", "amount"]:
        widget_key = f"claim_meta_{claim_key}_{field}"
        if widget_key in session:
            claim[field] = safe_str(session.get(widget_key, ""))

    dealer_key = f"claim_meta_{claim_key}_dealer"
    if dealer_key in session:
        claim["dealer"] = safe_str(session.get(dealer_key, ""))

    type_key = f"deduction_type_{claim_key}"
    amount_key = f"deduction_amount_{claim_key}"
    reason_key = f"deduction_reason_{claim_key}"
    if type_key in session:
        claim["deduction_type"] = safe_str(session.get(type_key, "none")) or "none"
    if amount_key in session:
        claim["deduction_amount"] = float(parse_amount_value(session.get(amount_key, 0)))
    if reason_key in session:
        claim["deduction_reason"] = safe_str(session.get(reason_key, ""))

    for check in ALL_CHECKS:
        evaluation = claim.setdefault("evaluations", {}).setdefault(check.key, new_evaluation(check))
        select_key = f"select_{claim_key}_{check.key}"
        if select_key in session:
            option = option_from_label(check, session.get(select_key))
            evaluation["option_key"] = option.key
            evaluation["status"] = option.status
            evaluation["points"] = option.points
            evaluation["max_points"] = check.max_points

        comment_key = f"comment_{claim_key}_{check.key}"
        if comment_key in session:
            evaluation["comment"] = safe_str(session.get(comment_key, ""))

    claim["claim_no"] = safe_str(claim.get("local_claim_no")) or safe_str(claim.get("hq_claim_no")) or safe_str(claim.get("claim_no"))
    refresh_claim_general_comment(claim)


def sync_all_claims_from_widget_state(claims: Dict[str, Dict[str, Any]]) -> None:
    for claim in claims.values():
        sync_claim_from_widget_state(claim)


def build_action_plan_rows(claims: Dict[str, Dict[str, Any]], language: str = "es") -> List[Dict[str, Any]]:
    sync_all_claims_from_widget_state(claims)
    rows = []
    for check in ALL_SCORING_CHECKS:
        entries = []
        total_lost = 0
        affected_claims = 0
        for claim in claims.values():
            evaluation = claim.get("evaluations", {}).get(check.key, {})
            points = evaluation.get("points")
            status = status_for_language(evaluation.get("status", ""), language) or ("Pendiente" if language == "es" else "Pending")
            comment = comment_for_language(evaluation_comment(claim, check), language)
            lost = check.max_points if points is None else max(0, check.max_points - int(points or 0))
            if lost > 0 or comment:
                affected_claims += 1
                total_lost += int(lost)
                claim_id = claim_identifier(claim, language)
                if comment:
                    entries.append(f"{claim_id}: {comment}")
                else:
                    if language == "es":
                        entries.append(f"{claim_id}: {status}, sin observación específica.")
                    else:
                        entries.append(f"{claim_id}: {status}, no specific observation.")
        if affected_claims:
            if language == "es":
                exception = f"{affected_claims} claim(s) con desviación u observación. Puntos perdidos: {total_lost}."
            else:
                exception = f"{affected_claims} claim(s) with deviation or observation. Lost points: {total_lost}."
            rows.append({
                "block": block_label(check.block_key, language),
                "parameter": check_label(check, language),
                "exception": exception,
                "observations": "\n".join(entries),
                "countermeasure": COUNTERMEASURES.get(check.key, {}).get(language, TEXT[language]["generic_countermeasure"]),
                "lost_points": total_lost,
                "affected_claims": affected_claims,
            })
    # Ajustes económicos / deducciones: no forman parte de la puntuación, pero sí del plan de acción.
    deduction_entries = []
    total_deduction = 0.0
    affected_deductions = 0
    for claim in claims.values():
        dtype = safe_str(claim.get("deduction_type", "none")) or "none"
        amount = claim_deduction_amount(claim)
        reason = claim_deduction_reason(claim, language)
        if dtype != "none" or amount > 0 or reason:
            affected_deductions += 1
            total_deduction += amount
            claim_id = claim_identifier(claim, language)
            label = deduction_type_label(dtype, language)
            amount_text = f"{amount:.2f}" if amount > 0 else ("importe completo" if language == "es" and dtype == "total" else "full amount" if dtype == "total" else "0")
            if reason:
                deduction_entries.append(f"{claim_id}: {label} - {amount_text}. {reason}")
            else:
                deduction_entries.append(f"{claim_id}: {label} - {amount_text}.")
    if affected_deductions:
        if language == "es":
            exception = f"{affected_deductions} claim(s) con deducción registrada. Importe deducido informado: {total_deduction:.2f}."
            countermeasure = "Revisar el importe reclamado y comunicar claramente al dealer la deducción aplicada y su motivo."
            parameter = "Deducción / ajuste económico"
            block = "Coste de garantía"
        else:
            exception = f"{affected_deductions} claim(s) with registered deduction. Reported deducted amount: {total_deduction:.2f}."
            countermeasure = "Review the claimed amount and clearly communicate the applied deduction and its reason to the dealer."
            parameter = "Deduction / financial adjustment"
            block = "Warranty cost"
        rows.append({
            "block": block,
            "parameter": parameter,
            "exception": exception,
            "observations": "\n".join(deduction_entries),
            "countermeasure": countermeasure,
            "lost_points": 0,
            "affected_claims": affected_deductions,
        })
    return rows


def ai_action_plan_rows_for_language(language: str = "es") -> List[Dict[str, str]]:
    """Filas estructuradas devueltas por Gemini para la hoja Improvement/action plan."""
    key = "ai_action_plan_rows_es" if language == "es" else "ai_action_plan_rows_en"
    try:
        raw_rows = st.session_state.get(key, []) or []
    except Exception:
        raw_rows = []
    if not isinstance(raw_rows, list):
        return []
    rows: List[Dict[str, str]] = []
    for raw in raw_rows:
        if not isinstance(raw, dict):
            continue
        rows.append({
            "group": safe_str(raw.get("group", "")),
            "parameter": safe_str(raw.get("parameter", "")),
            "evaluation": "X" if safe_str(raw.get("evaluation", "")).upper() == "X" else "V",
            "observations": safe_str(raw.get("observations", "")),
            "countermeasures": safe_str(raw.get("countermeasures", "")),
        })
    return rows


def build_action_plan_scorecard_rows(claims: Dict[str, Dict[str, Any]], language: str = "es") -> List[Dict[str, str]]:
    """
    Construye la página 3 en formato checklist/action plan.
    Siempre incluye todos los apartados y marca V/X.
    Si Gemini ha generado filas estructuradas, se usan sus observaciones/contramedidas,
    pero se fuerza el orden y los nombres exactos de la checklist.
    """
    sync_all_claims_from_widget_state(claims)

    ai_rows = ai_action_plan_rows_for_language(language)
    ai_by_parameter = {
        normalize_text(row.get("parameter", "")): row
        for row in ai_rows
        if row.get("parameter")
    }

    rows: List[Dict[str, str]] = []
    for check in ALL_SCORING_CHECKS:
        expected_parameter = check_label(check, language)
        group = "Claim" if check.block_key == "document" else "Old part"

        comments = []
        total_lost = 0
        affected_claims = 0
        for claim in claims.values():
            evaluation = claim.get("evaluations", {}).get(check.key, {})
            points = evaluation.get("points")
            lost = check.max_points if points is None else max(0, check.max_points - int(points or 0))
            comment = comment_for_language(evaluation_comment(claim, check), language)
            if lost > 0 or comment:
                affected_claims += 1
                total_lost += int(lost)
                if comment:
                    comments.append(comment)

        has_issue = affected_claims > 0 or total_lost > 0
        evaluation_mark = "X" if has_issue else "V"

        # Fallback automático si no hay Gemini o falta una fila concreta.
        if has_issue:
            unique_comments = []
            seen = set()
            for comment in comments:
                norm = normalize_text(comment)
                if norm and norm not in seen:
                    unique_comments.append(comment)
                    seen.add(norm)
            if language == "es":
                observation = f"Se detectan desviaciones u observaciones en {affected_claims} claim(s)."
                if total_lost:
                    observation += f" Puntos perdidos: {total_lost}."
                if unique_comments:
                    observation += " " + " / ".join(unique_comments[:3])
            else:
                observation = f"Deviations or observations were identified in {affected_claims} claim(s)."
                if total_lost:
                    observation += f" Lost points: {total_lost}."
                if unique_comments:
                    observation += " " + " / ".join(unique_comments[:3])
            countermeasure = COUNTERMEASURES.get(check.key, {}).get(language, TEXT[language]["generic_countermeasure"])
        else:
            observation = ""
            countermeasure = ""

        ai_row = ai_by_parameter.get(normalize_text(expected_parameter), {})
        if ai_row:
            # El nombre y el orden se fuerzan desde la checklist; solo aprovechamos contenido IA.
            ai_evaluation = safe_str(ai_row.get("evaluation", "")).upper()
            evaluation_mark = "X" if ai_evaluation == "X" or has_issue else "V"
            if evaluation_mark == "X":
                observation = safe_str(ai_row.get("observations", "")) or observation
                countermeasure = safe_str(ai_row.get("countermeasures", "")) or countermeasure
            else:
                observation = ""
                countermeasure = ""

        rows.append({
            "group": group,
            "parameter": expected_parameter,
            "evaluation": evaluation_mark,
            "observations": observation,
            "countermeasures": countermeasure,
        })

    return rows


# =============================================================================
# CARGA RÁPIDA DE CLAIMS DESDE TABLA EDITABLE
# =============================================================================

CLAIM_PASTE_COLUMNS = ["Claim HQ / IDMS", "Claim España / Dealer", "Dealer"]


def blank_claim_paste_dataframe(rows: int = 12) -> pd.DataFrame:
    """Crea una tabla vacía para pegar claims desde Excel."""
    return pd.DataFrame(
        [
            {"Claim HQ / IDMS": "", "Claim España / Dealer": "", "Dealer": ""}
            for _ in range(rows)
        ]
    )


def read_claims_from_pasted_table(table_df: pd.DataFrame) -> Dict[str, Dict[str, Any]]:
    """
    Convierte la tabla editable de Streamlit en claims.

    Formato esperado:
    - Columna 1: Claim HQ / IDMS (2810..., TAC..., etc.)
    - Columna 2: Claim España / Dealer (CO...)
    - Columna 3: Dealer / instalación

    Si una fila solo trae un identificador, se intenta inferir si es local o HQ.
    Si Dealer viene vacío, se aplicará el dealer general de la barra lateral.
    """
    claims: Dict[str, Dict[str, Any]] = {}

    if table_df is None or table_df.empty:
        return claims

    for _, row in table_df.iterrows():
        local_value = safe_str(row.get("Claim España / Dealer", ""))
        hq_value = safe_str(row.get("Claim HQ / IDMS", ""))
        dealer_value = safe_str(row.get("Dealer", ""))

        # Si el usuario pega un único identificador en la primera columna, inferimos.
        if local_value and not hq_value:
            inferred_local, inferred_hq = infer_claim_ids(local_value)
            local_value = inferred_local or local_value
            hq_value = inferred_hq

        # Si el usuario pega un único identificador en la segunda columna, inferimos también.
        if hq_value and not local_value:
            inferred_local, inferred_hq = infer_claim_ids(hq_value)
            local_value = inferred_local
            hq_value = inferred_hq or hq_value

        if not local_value and not hq_value:
            continue

        claim = upsert_claim(
            claims,
            claim_no=local_value or hq_value,
            local_claim_no=local_value,
            hq_claim_no=hq_value,
        )
        claim["local_claim_no"] = local_value
        claim["hq_claim_no"] = hq_value
        claim["claim_no"] = local_value or hq_value
        if dealer_value:
            claim["dealer"] = dealer_value

    return claims


# =============================================================================
# DATAFRAMES / EXPORTACIÓN
# =============================================================================


def build_summary_dataframe(claims: Dict[str, Dict[str, Any]], language: str = "es") -> pd.DataFrame:
    sync_all_claims_from_widget_state(claims)
    t = TEXT[language]
    rows = []
    for claim in claims.values():
        score = calculate_claim_score(claim)
        rows.append({
            t["claim_no"]: claim_identifier(claim, language),
            t["other_id"]: claim_other_identifier(claim, language),
            t["dealer"]: claim.get("dealer", ""),
            t["vin"]: claim.get("vin", ""),
            t["model"]: claim.get("model", ""),
            t["amount"]: claim.get("amount", ""),
            t["deduction_type"]: deduction_type_label(claim.get("deduction_type", "none"), language),
            t["deduction_amount"]: claim_deduction_amount(claim),
            t["deduction_reason"]: claim_deduction_reason(claim, language),
            t["doc_score"]: score["doc_points"],
            t["old_score"]: score["old_points"],
            t["total_score"]: score["total_points"],
            t["success"]: score["success_percent"],
            t["result"]: result_label(score["total_points"], language),
            t["pending"]: " | ".join(pending_items_for_language(claim, language)),
            t["comments"]: general_comment_for_export(claim, language),
        })
    return pd.DataFrame(rows)


def build_detail_dataframe(claims: Dict[str, Dict[str, Any]], language: str = "es") -> pd.DataFrame:
    sync_all_claims_from_widget_state(claims)
    t = TEXT[language]
    rows = []
    for claim in claims.values():
        for check in ALL_CHECKS:
            evaluation = claim.get("evaluations", {}).get(check.key, {})
            points = evaluation.get("points")
            max_points = check.max_points
            rows.append({
                t["claim_no"]: claim_identifier(claim, language),
                t["other_id"]: claim_other_identifier(claim, language),
                t["dealer"]: claim.get("dealer", ""),
                t["vin"]: claim.get("vin", ""),
                t["model"]: claim.get("model", ""),
                t["amount"]: claim.get("amount", ""),
                t["deduction_type"]: deduction_type_label(claim.get("deduction_type", "none"), language),
                t["deduction_amount"]: claim_deduction_amount(claim),
                t["deduction_reason"]: claim_deduction_reason(claim, language),
                "Bloque" if language == "es" else "Section": block_label(check.block_key, language),
                "Apartado" if language == "es" else "Item": check_label(check, language),
                "Estado" if language == "es" else "Status": status_for_language(evaluation.get("status", ""), language),
                "Puntos" if language == "es" else "Points": points,
                "Máximo" if language == "es" else "Max": max_points,
                "Pérdida" if language == "es" else "Lost": "" if points is None else max_points - int(points or 0),
                "Comentario apartado" if language == "es" else "Item observation": comment_for_language(evaluation.get("comment", ""), language),
                "Criterio" if language == "es" else "Criterion": check_guidance(check, language),
            })
    return pd.DataFrame(rows)


def export_analytical_excel(claims: Dict[str, Dict[str, Any]], audit_name: str, dealer: str, auditor: str, language: str = "es", audit_date_value: Any = None) -> bytes:
    output = BytesIO()
    t = TEXT[language]
    audit_score = calculate_audit_score(claims)
    summary_df = build_summary_dataframe(claims, language)
    detail_df = build_detail_dataframe(claims, language)
    action_df = pd.DataFrame(build_action_plan_rows(claims, language))
    if action_df.empty:
        action_df = pd.DataFrame(columns=["block", "parameter", "exception", "observations", "countermeasure", "lost_points", "affected_claims"])

    cover_df = pd.DataFrame([
        {"Campo" if language == "es" else "Field": t["report_title"], "Valor" if language == "es" else "Value": audit_name},
        {"Campo" if language == "es" else "Field": t["dealer"], "Valor" if language == "es" else "Value": dealer},
        {"Campo" if language == "es" else "Field": "Auditor", "Valor" if language == "es" else "Value": auditor},
        {"Campo" if language == "es" else "Field": "Fecha auditoría" if language == "es" else "Audit date", "Valor" if language == "es" else "Value": audit_date_to_text(audit_date_value, language)},
        {"Campo" if language == "es" else "Field": "Fecha exportación" if language == "es" else "Export date", "Valor" if language == "es" else "Value": datetime.now().strftime("%Y-%m-%d %H:%M")},
        {"Campo" if language == "es" else "Field": "Resultado", "Valor" if language == "es" else "Value": f"{audit_score['success_percent']:.1f}% ({audit_score['total_points']}/{audit_score['max_points']})"},
        {"Campo" if language == "es" else "Field": "Deducción total registrada" if language == "es" else "Total registered deduction", "Valor" if language == "es" else "Value": f"{sum(claim_deduction_amount(c) for c in claims.values()):.2f}"},
    ])

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        cover_df.to_excel(writer, index=False, sheet_name="Resumen" if language == "es" else "Summary")
        summary_export = summary_df.copy()
        if t["success"] in summary_export.columns:
            summary_export[t["success"]] = pd.to_numeric(summary_export[t["success"]], errors="coerce") / 100
        summary_export.to_excel(writer, index=False, sheet_name="Claims")
        detail_df.to_excel(writer, index=False, sheet_name="Detalle" if language == "es" else "Detail")
        action_df.rename(columns={
            "block": "Bloque" if language == "es" else "Section",
            "parameter": t["auditable_list"],
            "exception": t["exception_comments"],
            "observations": t["observations"],
            "countermeasure": t["countermeasure"],
            "lost_points": "Puntos perdidos" if language == "es" else "Lost points",
            "affected_claims": "Claims afectadas" if language == "es" else "Affected claims",
        }).to_excel(writer, index=False, sheet_name="Plan de mejora" if language == "es" else "Action plan")

        workbook = writer.book
        header_format = workbook.add_format({"bold": True, "font_color": "white", "bg_color": "#1F4E78", "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True})
        body_format = workbook.add_format({"border": 1, "valign": "top", "text_wrap": True})
        percent_format = workbook.add_format({"num_format": "0.0%", "border": 1, "valign": "top"})
        for sheet_name, worksheet in writer.sheets.items():
            df = {"Resumen" if language == "es" else "Summary": cover_df, "Claims": summary_export, "Detalle" if language == "es" else "Detail": detail_df, "Plan de mejora" if language == "es" else "Action plan": action_df}[sheet_name]
            rows, cols = df.shape
            worksheet.freeze_panes(1, 0)
            if cols:
                worksheet.autofilter(0, 0, max(rows, 1), cols - 1)
            for col_num, value in enumerate(df.columns.values):
                worksheet.write(0, col_num, value, header_format)
                worksheet.set_column(col_num, col_num, 22 if col_num < 5 else 35, body_format)
            if sheet_name == "Claims" and t["success"] in summary_export.columns:
                percent_col = summary_export.columns.get_loc(t["success"])
                worksheet.set_column(percent_col, percent_col, 14, percent_format)
    return output.getvalue()


def export_scorecard_excel(claims: Dict[str, Dict[str, Any]], audit_name: str, dealer: str, auditor: str, language: str = "es", audit_date_value: Any = None) -> bytes:
    output = BytesIO()
    workbook = None
    t = TEXT[language]
    try:
        import xlsxwriter  # noqa: F401
        workbook = xlsxwriter.Workbook(output, {"in_memory": True})

        fmt_title = workbook.add_format({"bold": True, "font_size": 14, "font_color": "#1F4E78", "valign": "vcenter", "text_wrap": True})
        fmt_header = workbook.add_format({"bold": True, "font_color": "white", "bg_color": "#1F4E78", "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True})
        fmt_body = workbook.add_format({"border": 1, "valign": "top", "text_wrap": True})
        fmt_bold = workbook.add_format({"bold": True, "border": 1, "valign": "top", "text_wrap": True})
        fmt_int = workbook.add_format({"border": 1, "valign": "top", "num_format": "0"})
        fmt_percent = workbook.add_format({"border": 1, "align": "center", "valign": "vcenter", "num_format": "0.0%", "bold": True})
        fmt_formula_percent = workbook.add_format({"border": 1, "align": "center", "valign": "vcenter", "num_format": "0.0%"})
        fmt_ok = workbook.add_format({"border": 1, "align": "center", "valign": "vcenter", "bg_color": "#E2F0D9", "bold": True})
        fmt_mid = workbook.add_format({"border": 1, "align": "center", "valign": "vcenter", "bg_color": "#FFF2CC", "bold": True})
        fmt_bad = workbook.add_format({"border": 1, "align": "center", "valign": "vcenter", "bg_color": "#F4CCCC", "bold": True})
        fmt_gray = workbook.add_format({"border": 1, "align": "center", "valign": "vcenter", "bg_color": "#E7E6E6"})

        def write_row(ws, row_idx, values, fmt=fmt_body):
            for col_idx, value in enumerate(values):
                ws.write(row_idx, col_idx, value, fmt)

        def write_headers(ws, row_idx, headers):
            ws.set_row(row_idx, 30)
            for col_idx, value in enumerate(headers):
                ws.write(row_idx, col_idx, value, fmt_header)

        def set_widths(ws, widths):
            for col_idx, width in enumerate(widths):
                ws.set_column(col_idx, col_idx, width)

        def result_format(points):
            try:
                p = float(points)
            except Exception:
                return fmt_gray
            if p >= 80:
                return fmt_ok
            if p >= 60:
                return fmt_mid
            return fmt_bad

        def evaluation_value(claim, check):
            points = claim["evaluations"].get(check.key, {}).get("points")
            return "" if points is None else points

        # ------------------------------------------------------------------
        # Scorecard / Boletín
        # ------------------------------------------------------------------
        grade_ws = workbook.add_worksheet(t["scorecard"][:31])
        grade_ws.merge_range("A1:P1", f"{t['scorecard']} - Warranty Audit", fmt_title)
        grade_headers = [
            t["claim_no"], t["other_id"], t["dealer"], t["vin"], t["model"], t["amount"],
            t["deduction_type"], t["deduction_amount"], t["deduction_reason"],
            t["doc_score"], t["old_score"], t["total_score"], t["success"], t["result"], t["pending"], t["comments"],
        ]
        write_headers(grade_ws, 2, grade_headers)
        set_widths(grade_ws, [20, 20, 24, 22, 18, 14, 18, 14, 45, 18, 18, 14, 14, 16, 45, 70])
        grade_ws.freeze_panes(3, 0)
        for idx, claim in enumerate(claims.values(), start=3):
            score = calculate_claim_score(claim)
            row_values = [
                claim_identifier(claim, language),
                claim_other_identifier(claim, language),
                claim.get("dealer", dealer or ""),
                claim.get("vin", ""),
                claim.get("model", ""),
                claim.get("amount", ""),
                deduction_type_label(claim.get("deduction_type", "none"), language),
                claim_deduction_amount(claim),
                claim_deduction_reason(claim, language),
                score["doc_points"],
                score["old_points"],
                score["total_points"],
                score["success_percent"] / 100,
                result_label(score["total_points"], language),
                " | ".join(pending_items_for_language(claim, language)),
                general_comment_for_export(claim, language),
            ]
            write_row(grade_ws, idx, row_values, fmt_body)
            grade_ws.write_number(idx, 7, claim_deduction_amount(claim), fmt_int)
            grade_ws.write_number(idx, 9, score["doc_points"], fmt_int)
            grade_ws.write_number(idx, 10, score["old_points"], fmt_int)
            grade_ws.write_number(idx, 11, score["total_points"], result_format(score["total_points"]))
            grade_ws.write_number(idx, 12, score["success_percent"] / 100, fmt_percent)
            grade_ws.write(idx, 13, result_label(score["total_points"], language), result_format(score["total_points"]))
        if claims:
            grade_ws.autofilter(2, 0, len(claims) + 2, len(grade_headers) - 1)
            grade_ws.conditional_format(3, 12, len(claims) + 2, 12, {"type": "3_color_scale", "min_color": "#F4CCCC", "mid_color": "#FFF2CC", "max_color": "#E2F0D9"})

        # ------------------------------------------------------------------
        # Content
        # ------------------------------------------------------------------
        content_ws = workbook.add_worksheet("content")
        content_ws.merge_range("A1:D1", f"Warranty audit report card / {t['scorecard']}", fmt_title)
        content_rows = [
            ["Audit name" if language == "en" else "Auditoría", audit_name or ""],
            [t["dealer"], dealer or ""],
            ["Auditor", auditor or ""],
            ["Audit date" if language == "en" else "Fecha auditoría", audit_date_to_text(audit_date_value, language)],
            ["Export date" if language == "en" else "Fecha exportación", datetime.now().strftime("%Y-%m-%d %H:%M")],
            ["Scoring rule" if language == "en" else "Regla de puntuación", "Claim document checklist I = 58 / Claim old parts checklist II = 42 / Total = 100"],
            ["N/A rule" if language == "en" else "Regla N/A", "No aplica = maximum score" if language == "en" else "No aplica = puntuación máxima del apartado"],
            ["Manual observations" if language == "en" else "Observaciones manuales", t["not_translated_note"]],
            ["Deduction" if language == "en" else "Deducción", "Financial deduction is informative and does not modify the audit score unless the claim is marked as not covered by warranty." if language == "en" else "La deducción económica es informativa y no modifica la nota salvo que se marque la claim como no cubierta por garantía."],
        ]
        set_widths(content_ws, [28, 90, 18, 18])
        for row_idx, row in enumerate(content_rows, start=2):
            content_ws.write(row_idx, 0, row[0], fmt_bold)
            content_ws.write(row_idx, 1, row[1], fmt_body)

        audit_score = calculate_audit_score(claims)
        content_ws.write(11, 0, "Global result" if language == "en" else "Resultado global", fmt_bold)
        content_ws.write(11, 1, f"{audit_score['success_percent']:.1f}% ({audit_score['total_points']}/{audit_score['max_points']})", result_format(audit_score["success_percent"]))

        # ------------------------------------------------------------------
        # Claim document checklist I
        # ------------------------------------------------------------------
        doc_ws = workbook.add_worksheet("Claim document checklist I")
        doc_headers = [
            "No.", t["claim_no"], t["other_id"],
            *[check_label(check, language) for check in DOCUMENT_CHECKS],
            t["doc_score"], "Document %" if language == "en" else "Resultado documentación %",
            t["deduction_amount"], t["deduction_reason"], t["comments"],
        ]
        write_headers(doc_ws, 0, doc_headers)
        set_widths(doc_ws, [8, 20, 20, 12, 20, 18, 12, 16, 12, 18, 22, 12, 18, 18, 14, 45, 60])
        doc_ws.freeze_panes(1, 3)
        for idx, claim in enumerate(claims.values(), start=1):
            excel_row = idx + 1
            row = [idx, claim_identifier(claim, language), claim_other_identifier(claim, language)]
            row.extend(evaluation_value(claim, check) for check in DOCUMENT_CHECKS)
            row.append(f"=SUM(D{excel_row}:L{excel_row})")
            row.append(f"=M{excel_row}/{MAX_DOCUMENT_POINTS}")
            row.append(claim_deduction_amount(claim))
            row.append(claim_deduction_reason(claim, language))
            row.append(general_comment_for_export(claim, language))
            write_row(doc_ws, idx, row, fmt_body)
            for col_idx in range(3, 12):
                if isinstance(row[col_idx], (int, float)):
                    doc_ws.write_number(idx, col_idx, row[col_idx], fmt_int)
            doc_ws.write_formula(idx, 12, f"=SUM(D{excel_row}:L{excel_row})", fmt_int)
            doc_ws.write_formula(idx, 13, f"=M{excel_row}/{MAX_DOCUMENT_POINTS}", fmt_formula_percent)
            doc_ws.write_number(idx, 14, claim_deduction_amount(claim), fmt_int)
        if claims:
            doc_ws.autofilter(0, 0, len(claims), len(doc_headers) - 1)
            doc_ws.conditional_format(1, 13, len(claims), 13, {"type": "3_color_scale", "min_color": "#F4CCCC", "mid_color": "#FFF2CC", "max_color": "#E2F0D9"})

        # ------------------------------------------------------------------
        # Claim old parts checklist II
        # ------------------------------------------------------------------
        old_ws = workbook.add_worksheet("Claim old parts checklist II")
        old_headers = [
            "No.", t["claim_no"], t["other_id"],
            *[check_label(check, language) for check in OLD_PARTS_CHECKS],
            t["old_score"], "Old parts %" if language == "en" else "Resultado piezas viejas %", t["deduction_amount"], t["deduction_reason"], t["comments"],
        ]
        write_headers(old_ws, 0, old_headers)
        set_widths(old_ws, [8, 20, 20, 22, 22, 18, 28, 24, 30, 18, 18, 14, 45, 65])
        old_ws.freeze_panes(1, 3)
        for idx, claim in enumerate(claims.values(), start=1):
            excel_row = idx + 1
            row = [idx, claim_identifier(claim, language), claim_other_identifier(claim, language)]
            row.extend(evaluation_value(claim, check) for check in OLD_PARTS_CHECKS)
            row.append(f"=SUM(D{excel_row}:I{excel_row})")
            row.append(f"=J{excel_row}/{MAX_OLD_PARTS_POINTS}")
            row.append(claim_deduction_amount(claim))
            row.append(claim_deduction_reason(claim, language))
            row.append(general_comment_for_export(claim, language))
            write_row(old_ws, idx, row, fmt_body)
            for col_idx in range(3, 9):
                if isinstance(row[col_idx], (int, float)):
                    old_ws.write_number(idx, col_idx, row[col_idx], fmt_int)
            old_ws.write_formula(idx, 9, f"=SUM(D{excel_row}:I{excel_row})", fmt_int)
            old_ws.write_formula(idx, 10, f"=J{excel_row}/{MAX_OLD_PARTS_POINTS}", fmt_formula_percent)
            old_ws.write_number(idx, 11, claim_deduction_amount(claim), fmt_int)
        if claims:
            old_ws.autofilter(0, 0, len(claims), len(old_headers) - 1)
            old_ws.conditional_format(1, 10, len(claims), 10, {"type": "3_color_scale", "min_color": "#F4CCCC", "mid_color": "#FFF2CC", "max_color": "#E2F0D9"})

        # ------------------------------------------------------------------
        # Improvement of claim issues III - página 3 / plan de acción
        # ------------------------------------------------------------------
        imp_ws = workbook.add_worksheet("Improvement of claim issues III")
        imp_ws.merge_range("A1:E1", "Action Plan" if language == "en" else "Plan de acción", fmt_title)
        set_widths(imp_ws, [14, 34, 14, 85, 85])

        imp_headers = [
            "",
            "Audited parameters" if language == "en" else "Parámetros auditados",
            "Evaluation" if language == "en" else "Evaluación",
            "Observations" if language == "en" else "Observaciones",
            "Countermeasures" if language == "en" else "Contramedidas",
        ]
        write_headers(imp_ws, 1, imp_headers)
        imp_ws.freeze_panes(2, 0)

        plan_rows = build_action_plan_scorecard_rows(claims, language)
        start_row = 2
        for row_idx, item in enumerate(plan_rows, start=start_row):
            write_row(imp_ws, row_idx, [
                item["group"],
                item["parameter"],
                item["evaluation"],
                item["observations"],
                item["countermeasures"],
            ], fmt_body)
            if item["evaluation"] == "X":
                imp_ws.write(row_idx, 2, "X", fmt_bad)
            else:
                imp_ws.write(row_idx, 2, "V", fmt_ok)

        # Agrupación visual tipo plantilla: Claim / Old part.
        if plan_rows:
            doc_count = len(DOCUMENT_CHECKS)
            old_count = len(OLD_PARTS_CHECKS)
            if doc_count > 1:
                imp_ws.merge_range(start_row, 0, start_row + doc_count - 1, 0, "Claim", fmt_body)
            if old_count > 1:
                imp_ws.merge_range(start_row + doc_count, 0, start_row + doc_count + old_count - 1, 0, "Old\npart", fmt_body)
            imp_ws.autofilter(1, 0, start_row + len(plan_rows) - 1, len(imp_headers) - 1)

        # ------------------------------------------------------------------
        # Evaluation content
        # ------------------------------------------------------------------
        eval_ws = workbook.add_worksheet("Evaluation content")
        eval_headers = ["Checklist", "Apartado" if language == "es" else "Item", "Máximo" if language == "es" else "Max", "Opciones" if language == "es" else "Options", "Criterio" if language == "es" else "Criterion"]
        write_headers(eval_ws, 0, eval_headers)
        set_widths(eval_ws, [30, 38, 12, 54, 95])
        eval_ws.freeze_panes(1, 0)
        for row_idx, check in enumerate(ALL_CHECKS, start=1):
            options_text = " | ".join(
                f"{option_display(option, language)}: {'-' if option.points is None else option.points}"
                for option in check.options
            )
            write_row(eval_ws, row_idx, [block_label(check.block_key, language), check_label(check, language), check.max_points, options_text, check_guidance(check, language)], fmt_body)
        eval_ws.autofilter(0, 0, len(ALL_CHECKS), len(eval_headers) - 1)

        workbook.close()
        workbook = None
        return output.getvalue()
    finally:
        if workbook is not None:
            workbook.close()


def generate_text_report(claims: Dict[str, Dict[str, Any]], audit_name: str, dealer: str, auditor: str, language: str = "es", audit_date_value: Any = None) -> str:
    ai_report = get_ai_report_for_language(language)
    if ai_report:
        return ai_report

    t = TEXT[language]
    audit_score = calculate_audit_score(claims)
    action_rows = build_action_plan_rows(claims, language)
    summary_df = build_summary_dataframe(claims, language)

    lines = []
    lines.append(f"{t['report_title']}: {audit_name or 'Warranty audit'}")
    lines.append(f"{t['dealer']}: {dealer or ('No informado' if language == 'es' else 'Not reported')}")
    lines.append(f"{'Fecha auditoría' if language == 'es' else 'Audit date'}: {audit_date_to_text(audit_date_value, language)}")
    lines.append(f"Auditor: {auditor or ('No informado' if language == 'es' else 'Not reported')}")
    lines.append("")
    lines.append(t["executive_summary"])
    if language == "es":
        lines.append(
            f"Se han revisado {audit_score['claims']} garantías, con {audit_score['completed_claims']} completadas. "
            f"El resultado global de la auditoría es {audit_score['success_percent']:.1f}% "
            f"({audit_score['total_points']}/{audit_score['max_points']} puntos)."
        )
    else:
        lines.append(
            f"{audit_score['claims']} claims have been reviewed, with {audit_score['completed_claims']} completed. "
            f"The global audit result is {audit_score['success_percent']:.1f}% "
            f"({audit_score['total_points']}/{audit_score['max_points']} points)."
        )
    lines.append("")
    lines.append(t["block_result"])
    if language == "es":
        lines.append(f"- Documentación de claim: {audit_score['doc_points']}/{audit_score['claims'] * MAX_DOCUMENT_POINTS} puntos.")
        lines.append(f"- Piezas viejas: {audit_score['old_points']}/{audit_score['claims'] * MAX_OLD_PARTS_POINTS} puntos.")
    else:
        lines.append(f"- Claim documentation: {audit_score['doc_points']}/{audit_score['claims'] * MAX_DOCUMENT_POINTS} points.")
        lines.append(f"- Old parts: {audit_score['old_points']}/{audit_score['claims'] * MAX_OLD_PARTS_POINTS} points.")
    lines.append("")
    lines.append(t["improvement_areas"])
    if not action_rows:
        lines.append(t["no_deviations"] + ".")
    else:
        for item in sorted(action_rows, key=lambda row: row["lost_points"], reverse=True)[:8]:
            lines.append(f"- {item['parameter']}: {item['exception']}")
    lines.append("")
    lines.append(t["lowest_claims"])
    if not summary_df.empty:
        score_col = t["total_score"]
        claim_col = t["claim_no"]
        comments_col = t["comments"]
        for _, row in summary_df.sort_values(score_col, ascending=True).head(5).iterrows():
            lines.append(f"- {row[claim_col]}: {row[score_col]}/100. {safe_str(row[comments_col])}")
    lines.append("")
    lines.append(t["conclusion"])
    if language == "es":
        lines.append("Se recomienda focalizar el plan de mejora en las áreas con mayor pérdida de puntos y reforzar la revisión previa de la documentación antes del envío de claims.")
    else:
        lines.append("It is recommended to focus the improvement plan on the areas with the highest lost points and reinforce documentation review before claim submission.")
    lines.append("")
    lines.append(t["not_translated_note"])
    return "\n".join(lines)


def clean_markdown_inline(text: Any) -> str:
    """Limpieza sencilla para que el Word no muestre marcas Markdown."""
    value = safe_str(text)
    value = re.sub(r"\*\*(.*?)\*\*", r"\1", value)
    value = re.sub(r"__(.*?)__", r"\1", value)
    value = re.sub(r"`([^`]*)`", r"\1", value)
    return value.strip()


def add_markdown_like_text_to_doc(doc, text: str, language: str = "es") -> None:
    """Convierte un texto Markdown/plain sencillo a párrafos y bullet points en Word."""
    heading_markers = {
        "es": ["resumen ejecutivo", "principales hallazgos", "principales areas", "principales áreas", "conclusion", "conclusión", "resultado por bloque", "claims con menor puntuacion", "claims con menor puntuación"],
        "en": ["executive summary", "main findings", "main improvement areas", "conclusion", "result by section", "lowest scoring claims"],
    }
    markers = heading_markers.get(language, heading_markers["es"])

    for raw_line in safe_str(text).splitlines():
        line = raw_line.strip()
        if not line:
            continue

        # Quitar títulos principales repetidos tipo "Informe de auditoría: ..." porque el Word ya lleva portada/cabecera.
        low = normalize_text(line)
        if low.startswith("informe de auditoria") or low.startswith("audit report") or low.startswith("dealer") or low.startswith("auditor") or low.startswith("fecha auditoria") or low.startswith("audit date"):
            continue

        if line.startswith("###"):
            doc.add_heading(clean_markdown_inline(line.lstrip("# ")), level=3)
        elif line.startswith("##"):
            doc.add_heading(clean_markdown_inline(line.lstrip("# ")), level=2)
        elif line.startswith("#"):
            doc.add_heading(clean_markdown_inline(line.lstrip("# ")), level=1)
        elif any(low == marker or low.startswith(marker + " ") for marker in markers):
            doc.add_heading(clean_markdown_inline(line.rstrip(":")), level=2)
        elif re.match(r"^[-*]\s+", line):
            doc.add_paragraph(clean_markdown_inline(re.sub(r"^[-*]\s+", "", line)), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            doc.add_paragraph(clean_markdown_inline(re.sub(r"^\d+[.)]\s+", "", line)), style="List Number")
        else:
            doc.add_paragraph(clean_markdown_inline(line))


def add_docx_table(doc, rows: List[List[Any]], header: bool = True, font_size: float = 9.0) -> None:
    if not rows:
        return
    from docx.shared import Pt

    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row_values):
            cell_text = clean_markdown_inline(value)
            cells[col_index].text = cell_text
            for paragraph in cells[col_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    if header and row_index == 0:
                        run.bold = True
    doc.add_paragraph()


def export_word_report_docx(claims: Dict[str, Dict[str, Any]], audit_name: str, dealer: str, auditor: str, language: str = "es", audit_date_value: Any = None) -> bytes:
    """Genera un informe Word presentable con resumen, plan de acción y firma."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except Exception as exc:
        raise RuntimeError("Para generar Word añade python-docx al requirements.txt") from exc

    sync_all_claims_from_widget_state(claims)
    t = TEXT[language]
    audit_score = calculate_audit_score(claims)
    report_text = generate_text_report(claims, audit_name, dealer, auditor, language, audit_date_value)
    action_rows = build_action_plan_scorecard_rows(claims, language)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            styles[style_name].font.name = "Aptos Display"
        except Exception:
            pass

    title_text = "Informe de auditoría de garantías" if language == "es" else "Warranty audit report"
    subtitle_text = audit_name or ("Auditoría garantías" if language == "es" else "Warranty audit")

    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(subtitle_text)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_rows = [
        [t["dealer"], dealer or ("No informado" if language == "es" else "Not reported")],
        ["Fecha auditoría" if language == "es" else "Audit date", audit_date_to_text(audit_date_value, language)],
        ["Auditor / firma" if language == "es" else "Auditor / signature", auditor or ("No informado" if language == "es" else "Not reported")],
        ["Resultado global" if language == "es" else "Global result", f"{audit_score['success_percent']:.1f}% ({audit_score['total_points']}/{audit_score['max_points']})"],
        [TEXT[language]["document_section"], f"{audit_score['doc_points']}/{audit_score['claims'] * MAX_DOCUMENT_POINTS}"],
        [TEXT[language]["old_parts_section"], f"{audit_score['old_points']}/{audit_score['claims'] * MAX_OLD_PARTS_POINTS}"],
    ]
    add_docx_table(doc, [["Campo" if language == "es" else "Field", "Valor" if language == "es" else "Value"], *meta_rows], font_size=9.5)

    doc.add_heading("Resumen ejecutivo" if language == "es" else "Executive summary", level=1)
    add_markdown_like_text_to_doc(doc, report_text, language)

    doc.add_page_break()
    doc.add_heading("Plan de acción" if language == "es" else "Action plan", level=1)
    if language == "es":
        headers = ["Grupo", "Parámetros auditados", "Evaluación", "Observaciones", "Contramedidas"]
    else:
        headers = ["Group", "Audited parameters", "Evaluation", "Observations", "Countermeasures"]
    plan_table_rows = [headers]
    for item in action_rows:
        plan_table_rows.append([
            item.get("group", ""),
            item.get("parameter", ""),
            item.get("evaluation", ""),
            item.get("observations", ""),
            item.get("countermeasures", ""),
        ])
    add_docx_table(doc, plan_table_rows, font_size=8.5)

    doc.add_paragraph()
    doc.add_heading("Firma" if language == "es" else "Signature", level=2)
    signature = doc.add_paragraph()
    signature.add_run(auditor or ("Auditor" if language == "en" else "Auditor")).bold = True
    signature.add_run("\n")
    signature.add_run(("NSC Warranty" if language == "en" else "NSC Warranty"))
    signature.add_run("\n")
    signature.add_run(("Fecha: " if language == "es" else "Date: ") + audit_date_to_text(audit_date_value, language))

    footer = section.footer.paragraphs[0]
    footer.text = f"{title_text} · {dealer or 'Dealer'} · {audit_date_to_text(audit_date_value, language)}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def export_bilingual_package_zip(claims: Dict[str, Dict[str, Any]], audit_name: str, dealer: str, auditor: str, audit_date_value: Any = None) -> bytes:
    output = BytesIO()
    base_name = build_audit_file_basename_from_date(dealer, auditor, audit_date_value)
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr(f"{base_name}/{base_name}.json", serialize_audit_workfile(claims, audit_name, dealer, auditor, audit_date_value))
        zip_file.writestr(f"{base_name}/{base_name}_ES_boletin.xlsx", export_scorecard_excel(claims, audit_name, dealer, auditor, "es", audit_date_value))
        zip_file.writestr(f"{base_name}/{base_name}_EN_scorecard.xlsx", export_scorecard_excel(claims, audit_name, dealer, auditor, "en", audit_date_value))
        zip_file.writestr(f"{base_name}/{base_name}_ES_informe.docx", export_word_report_docx(claims, audit_name, dealer, auditor, "es", audit_date_value))
        zip_file.writestr(f"{base_name}/{base_name}_EN_report.docx", export_word_report_docx(claims, audit_name, dealer, auditor, "en", audit_date_value))
        zip_file.writestr(f"{base_name}/{base_name}_ES_analitico.xlsx", export_analytical_excel(claims, audit_name, dealer, auditor, "es", audit_date_value))
        zip_file.writestr(f"{base_name}/{base_name}_EN_analytical.xlsx", export_analytical_excel(claims, audit_name, dealer, auditor, "en", audit_date_value))
    return output.getvalue()


# =============================================================================
# INTERFAZ STREAMLIT
# =============================================================================


def init_state():
    st.session_state.setdefault("claims", {})
    st.session_state.setdefault("selected_claim", None)
    st.session_state.setdefault("audit_name", "Auditoría garantías")
    st.session_state.setdefault("audit_dealer", "")
    st.session_state.setdefault("audit_auditor", "")
    st.session_state.setdefault("audit_date", datetime.now().date())
    st.session_state.setdefault("audit_section_index", 0)
    st.session_state.setdefault("claim_paste_editor_version", 0)


def clamp_index(index: Any, max_len: int) -> int:
    try:
        index = int(index)
    except Exception:
        index = 0
    if max_len <= 0:
        return 0
    return max(0, min(index, max_len - 1))


def set_audit_section_index(new_index: int, section_names: tuple) -> None:
    names = list(section_names)
    index = clamp_index(new_index, len(names))
    st.session_state.audit_section_index = index
    if names:
        # Esta key pertenece al radio, pero aquí se cambia dentro del callback,
        # antes de que el widget se vuelva a crear en el siguiente rerun.
        st.session_state.audit_section_radio = names[index]


def sync_audit_section_from_radio(section_names: tuple) -> None:
    names = list(section_names)
    selected = st.session_state.get("audit_section_radio", names[0] if names else "")
    if selected in names:
        st.session_state.audit_section_index = names.index(selected)
    else:
        set_audit_section_index(0, section_names)


def set_selected_claim_keep_section(claim_key: str, section_names: tuple) -> None:
    """
    Cambia de claim manteniendo el apartado activo.
    Esto encaja con el flujo real de auditoría:
    primero se revisa documentación para todas las claims y después piezas viejas en almacén.
    """
    st.session_state.selected_claim = claim_key
    # No tocamos audit_section_index: el siguiente rerun mantiene I. Documentación,
    # II. Piezas viejas, comentarios o informe según donde estuvieras.


def set_selected_claim_and_section(claim_key: str, section_index: int, section_names: tuple) -> None:
    """Cambia de claim y fuerza un apartado concreto."""
    st.session_state.selected_claim = claim_key
    set_audit_section_index(section_index, section_names)


def apply_default_dealer_to_blank_claims(claims: Dict[str, Dict[str, Any]], dealer: str) -> None:
    dealer = safe_str(dealer)
    if not dealer:
        return
    for claim in claims.values():
        if not safe_str(claim.get("dealer", "")):
            claim["dealer"] = dealer


def display_claim_option(key: str) -> str:
    claim = st.session_state.claims.get(key, {})
    local = safe_str(claim.get("local_claim_no", ""))
    hq = safe_str(claim.get("hq_claim_no", ""))
    score = calculate_claim_score(claim) if claim else {"total_points": 0}
    if local and hq:
        return f"{hq} / {local} · {score['total_points']}/100"
    return f"{hq or local or key} · {score['total_points']}/100"


def sync_claim_meta_field(claim: Dict[str, Any], field: str, label: str) -> str:
    claim_key = make_claim_key(claim.get("local_claim_no"), claim.get("hq_claim_no"), claim.get("claim_no"))
    widget_key = f"claim_meta_{claim_key}_{field}"
    if widget_key not in st.session_state:
        st.session_state[widget_key] = safe_str(claim.get(field, ""))
    value = st.text_input(label, key=widget_key)
    claim[field] = safe_str(value)
    return claim[field]


def render_claim_quick_card(claim: Dict[str, Any], default_dealer: str = ""):
    st.caption("Ficha de la garantía")
    key_before = make_claim_key(claim.get("local_claim_no"), claim.get("hq_claim_no"), claim.get("claim_no"))
    cols = st.columns(6)
    with cols[0]:
        hq = sync_claim_meta_field(claim, "hq_claim_no", "Claim HQ / IDMS")
    with cols[1]:
        local = sync_claim_meta_field(claim, "local_claim_no", "Claim ES / Dealer")
    with cols[2]:
        options = get_dealer_options(claim.get("dealer", "") or default_dealer)
        dealer_key = f"claim_meta_{key_before}_dealer"
        if not safe_str(claim.get("dealer", "")) and default_dealer:
            claim["dealer"] = default_dealer
        if dealer_key not in st.session_state or st.session_state[dealer_key] not in options:
            st.session_state[dealer_key] = claim.get("dealer", "") if claim.get("dealer", "") in options else ""
        dealer_value = st.selectbox("Dealer", options, key=dealer_key, format_func=format_dealer_option)
        claim["dealer"] = safe_str(dealer_value)
    with cols[3]:
        sync_claim_meta_field(claim, "vin", "VIN")
    with cols[4]:
        sync_claim_meta_field(claim, "model", "Modelo")
    with cols[5]:
        sync_claim_meta_field(claim, "amount", "Importe")

    # Actualizar claim_no interno y clave si cambian los IDs.
    claim["claim_no"] = local or hq or claim.get("claim_no", "")
    key_after = make_claim_key(local, hq, claim.get("claim_no", ""))
    if key_after and key_after != key_before and key_before in st.session_state.claims:
        st.session_state.claims[key_after] = st.session_state.claims.pop(key_before)
        st.session_state.selected_claim = key_after
        st.rerun()


def render_deduction_editor(claim: Dict[str, Any]):
    claim_key = make_claim_key(claim.get("local_claim_no"), claim.get("hq_claim_no"), claim.get("claim_no"))
    with st.expander("Deducción / ajuste económico", expanded=has_deduction(claim)):
        st.caption(
            "La deducción no cambia la nota de auditoría salvo que uses el botón 'No es garantía'. "
            "Para deducciones parciales, informa solo el importe y el motivo."
        )
        st.button(
            "🚫 No es garantía → 0/100 y deducción total",
            type="secondary",
            use_container_width=True,
            on_click=mark_claim_as_not_warranty_callback,
            args=(claim_key,),
            key=f"not_warranty_{claim_key}",
        )

        type_options = ["none", "partial", "total"]
        type_key = f"deduction_type_{claim_key}"
        if type_key not in st.session_state or st.session_state[type_key] not in type_options:
            st.session_state[type_key] = safe_str(claim.get("deduction_type", "none")) or "none"
        selected_type = st.selectbox(
            "Tipo de deducción",
            type_options,
            key=type_key,
            format_func=lambda value: deduction_type_label(value, "es"),
        )
        claim["deduction_type"] = selected_type

        amount_key = f"deduction_amount_{claim_key}"
        if amount_key not in st.session_state:
            st.session_state[amount_key] = claim_deduction_amount(claim)
        amount = st.number_input(
            "Importe a deducir",
            min_value=0.0,
            step=1.0,
            format="%.2f",
            key=amount_key,
        )
        claim["deduction_amount"] = float(amount or 0)

        reason_key = f"deduction_reason_{claim_key}"
        if reason_key not in st.session_state:
            st.session_state[reason_key] = safe_str(claim.get("deduction_reason", ""))
        reason = st.text_area("Motivo / comentario de deducción", key=reason_key, height=90)
        claim["deduction_reason"] = safe_str(reason)

        if selected_type == "none" and claim_deduction_amount(claim) > 0:
            st.warning("Hay importe de deducción informado, pero el tipo está en 'Sin deducción'.")
        elif selected_type == "partial" and claim_deduction_amount(claim) <= 0:
            st.info("Para deducción parcial, informa el importe a deducir.")
        elif selected_type == "total" and claim_deduction_amount(claim) <= 0:
            st.info("Deducción total seleccionada. Si conoces el importe de la claim, informa el importe a deducir.")


def render_check_editor(claim: Dict[str, Any], checks: List[AuditCheck]):
    claim_key = make_claim_key(claim.get("local_claim_no"), claim.get("hq_claim_no"), claim.get("claim_no"))
    for check in checks:
        evaluation = claim.setdefault("evaluations", {}).setdefault(check.key, new_evaluation(check))
        current_option = option_from_key(check, safe_str(evaluation.get("option_key", "ok")))
        labels = option_labels(check, "es")
        current_label = option_display(current_option, "es")
        current_index = labels.index(current_label) if current_label in labels else 0
        with st.container(border=True):
            cols = st.columns([2.2, 1.2, 2.8])
            with cols[0]:
                st.markdown(f"**{check.label_es}**")
                st.caption(f"Máximo: {check.max_points} puntos · {check.guidance_es}")
            with cols[1]:
                selected = st.selectbox("Evaluación", labels, index=current_index, key=f"select_{claim_key}_{check.key}", label_visibility="collapsed")
                option = option_from_label(check, selected)
                evaluation["option_key"] = option.key
                evaluation["status"] = option.status
                evaluation["points"] = option.points
                evaluation["max_points"] = check.max_points
                if option.points is None:
                    st.warning("Pendiente")
                else:
                    st.metric("Puntos", f"{option.points}/{check.max_points}")
            with cols[2]:
                comment_key = f"comment_{claim_key}_{check.key}"
                if comment_key not in st.session_state:
                    st.session_state[comment_key] = evaluation.get("comment", "")
                evaluation["comment"] = st.text_area("Observación del apartado", key=comment_key, height=85)

    refresh_claim_general_comment(claim)


def render_comments_editor(claim: Dict[str, Any]):
    st.caption(
        "El comentario general se actualiza automáticamente con las observaciones "
        "que escribas en los apartados. Para cambiarlo, cambia la observación del apartado correspondiente."
    )

    sync_claim_from_widget_state(claim)
    comment_es = refresh_claim_general_comment(claim)
    comment_en = auto_general_comment(claim, "en")

    if not comment_es:
        st.info("Todavía no hay observaciones escritas en los apartados ni motivo de deducción.")

    tab_es, tab_en = st.tabs(["Comentario automático ES", "Automatic comment EN"])
    with tab_es:
        st.text_area(
            "Comentario general automático de la claim",
            value=comment_es,
            height=180,
            disabled=True,
        )
    with tab_en:
        st.text_area(
            "Automatic general claim comment",
            value=comment_en,
            height=180,
            disabled=True,
        )


def render_report_section(claims: Dict[str, Dict[str, Any]], audit_name: str, dealer: str, auditor: str, base_name: str, audit_date_value: Any = None):
    st.subheader("Informe y plan de acción")

    with st.expander("✨ Generación con Gemini", expanded=not bool(get_ai_report_for_language("es"))):
        st.caption(
            "Gemini no cambia las puntuaciones. Solo redacta el informe ejecutivo, "
            "resume las observaciones, traduce al inglés y propone un plan de acción con la estructura de la checklist."
        )
        secret_key = get_secret_value("GEMINI_API_KEY", "GOOGLE_API_KEY")
        if secret_key:
            st.success("API key detectada en Streamlit Secrets / variable de entorno.")
            api_key = secret_key
        else:
            api_key = st.text_input(
                "Gemini API key temporal",
                type="password",
                help="Para Streamlit Cloud, mejor guárdala en Secrets como GEMINI_API_KEY. Si la pegas aquí, solo se usa en esta sesión.",
                key="gemini_api_key_input",
            )

        ai_cols = st.columns([1, 1, 2])
        with ai_cols[0]:
            if st.button("Generar con Gemini", type="primary", use_container_width=True):
                ok, message = generate_gemini_audit_outputs(
                    claims,
                    audit_name,
                    dealer,
                    auditor,
                    api_key,
                    GEMINI_MODEL_DEFAULT,
                )
                if ok:
                    st.success(message)
                    st.rerun()
                else:
                    st.error(message)
        with ai_cols[1]:
            if st.button("Borrar texto IA", use_container_width=True):
                clear_ai_outputs()
                st.rerun()
        with ai_cols[2]:
            generated_at = get_ai_output("ai_generated_at")
            if generated_at:
                st.caption(f"Último informe IA: {generated_at}")
            else:
                st.caption("Si no generas con Gemini, se usa el informe básico automático.")

    tabs = st.tabs(["Informe español", "Report English", "Plan acción ES", "Action plan EN"])
    with tabs[0]:
        report_es = generate_text_report(claims, audit_name, dealer, auditor, "es", audit_date_value)
        st.markdown(report_es)
        st.download_button("Descargar informe ES .docx", data=export_word_report_docx(claims, audit_name, dealer, auditor, "es", audit_date_value), file_name=f"{base_name}_ES_informe.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with tabs[1]:
        report_en = generate_text_report(claims, audit_name, dealer, auditor, "en", audit_date_value)
        st.markdown(report_en)
        st.download_button("Download EN report .docx", data=export_word_report_docx(claims, audit_name, dealer, auditor, "en", audit_date_value), file_name=f"{base_name}_EN_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with tabs[2]:
        action_rows_es = build_action_plan_scorecard_rows(claims, "es")
        st.dataframe(pd.DataFrame(action_rows_es).rename(columns={
            "group": "Grupo",
            "parameter": "Parámetros auditados",
            "evaluation": "Evaluación",
            "observations": "Observaciones",
            "countermeasures": "Contramedidas",
        }), use_container_width=True, hide_index=True)
        action_es = get_ai_action_plan_for_language("es")
        if action_es:
            with st.expander("Ver tabla generada por Gemini en texto", expanded=False):
                st.markdown(action_es)
    with tabs[3]:
        action_rows_en = build_action_plan_scorecard_rows(claims, "en")
        st.dataframe(pd.DataFrame(action_rows_en).rename(columns={
            "group": "Group",
            "parameter": "Audited parameters",
            "evaluation": "Evaluation",
            "observations": "Observations",
            "countermeasures": "Countermeasures",
        }), use_container_width=True, hide_index=True)
        action_en = get_ai_action_plan_for_language("en")
        if action_en:
            with st.expander("View Gemini-generated table as text", expanded=False):
                st.markdown(action_en)


def render_claim_paste_loader(dealer: str, show_title: bool = True):
    """
    Muestra una tabla editable tipo Excel para pegar claims directamente en la zona principal.
    Permite pegar claim HQ, claim española y dealer/instalación.
    """
    if show_title:
        st.subheader("Pegar lista de claims")

    st.caption(
        "Copia tres columnas desde Excel y pégalas aquí: "
        "Claim HQ/IDMS, Claim España/Dealer y Dealer/instalación."
    )

    editor_key = f"claims_paste_editor_{st.session_state.claim_paste_editor_version}"
    pasted_claims_df = st.data_editor(
        blank_claim_paste_dataframe(14),
        key=editor_key,
        num_rows="dynamic",
        use_container_width=True,
        hide_index=True,
        column_config={
            "Claim HQ / IDMS": st.column_config.TextColumn(
                "Claim HQ / IDMS",
                help="Identificador HQ/IDMS para scorecard en inglés, por ejemplo 2810... o TAC...",
            ),
            "Claim España / Dealer": st.column_config.TextColumn(
                "Claim España / Dealer",
                help="Identificador local para boletín al dealer, por ejemplo CO2026...",
            ),
            "Dealer": st.column_config.TextColumn(
                "Dealer",
                help="Dealer o instalación concreta. Si lo dejas vacío, se usará el dealer general de la barra lateral.",
            ),
        },
    )

    append_to_existing = False
    if st.session_state.claims:
        append_to_existing = st.checkbox(
            "Añadir a las claims ya cargadas",
            value=False,
            key="append_pasted_claims_main",
        )

    claim_table_cols = st.columns([1, 1, 3])
    with claim_table_cols[0]:
        if st.button("Cargar tabla", type="primary", use_container_width=True, key="load_claim_table_main"):
            try:
                loaded_claims = read_claims_from_pasted_table(pasted_claims_df)
                if not loaded_claims:
                    st.warning("La tabla está vacía. Pega al menos una claim española o de HQ.")
                else:
                    if dealer:
                        apply_default_dealer_to_blank_claims(loaded_claims, dealer)

                    if append_to_existing and st.session_state.claims:
                        st.session_state.claims.update(loaded_claims)
                    else:
                        st.session_state.claims = loaded_claims

                    st.session_state.selected_claim = next(iter(loaded_claims.keys()))
                    st.success(f"Cargadas {len(loaded_claims)} claims desde la tabla, con dealer por instalación si venía informado.")
                    st.rerun()
            except Exception as exc:
                st.error(f"No se pudo cargar la tabla: {exc}")

    with claim_table_cols[1]:
        if st.button("Limpiar tabla", use_container_width=True, key="clear_claim_table_main"):
            st.session_state.claim_paste_editor_version += 1
            st.rerun()

    with claim_table_cols[2]:
        st.caption("Tip: pega tres columnas desde Excel: HQ/IDMS · España/Dealer · Dealer/instalación.")

def main():
    st.set_page_config(page_title="Warranty Audit Assistant", page_icon="🧾", layout="wide")
    init_state()

    st.title("🧾 Warranty Audit Assistant")
    st.caption("Herramienta interna para revisar claims, calcular puntuación y generar automáticamente boletín/plan de acción en español e inglés.")

    with st.sidebar:
        st.header("Auditoría")
        st.session_state.audit_name = st.text_input("Nombre auditoría", value=st.session_state.audit_name)
        st.session_state.audit_date = st.date_input(
            "Fecha auditoría",
            value=parse_audit_date(st.session_state.get("audit_date", datetime.now().date())),
        )
        dealer_options = get_dealer_options(st.session_state.audit_dealer)
        if st.session_state.audit_dealer not in dealer_options:
            dealer_options.insert(1, st.session_state.audit_dealer)
        st.session_state.audit_dealer = st.selectbox("Dealer general / grupo", dealer_options, index=dealer_options.index(st.session_state.audit_dealer) if st.session_state.audit_dealer in dealer_options else 0, format_func=format_dealer_option)
        st.session_state.audit_auditor = st.text_input("Auditor", value=st.session_state.audit_auditor)
        dealer = safe_str(st.session_state.audit_dealer)
        auditor = safe_str(st.session_state.audit_auditor)
        audit_name = safe_str(st.session_state.audit_name)
        audit_date_value = parse_audit_date(st.session_state.get("audit_date", datetime.now().date()))
        base_name = build_audit_file_basename_from_date(dealer, auditor, audit_date_value)

        st.divider()
        st.subheader("Cargar / continuar")
        workfile = st.file_uploader("Cargar auditoría guardada (.json)", type=["json"], key="workfile_upload")
        if workfile is not None and st.button("Cargar JSON"):
            try:
                claims, loaded_name, loaded_dealer, loaded_auditor, loaded_date = load_audit_workfile(workfile)
                st.session_state.claims = claims
                st.session_state.audit_name = loaded_name
                st.session_state.audit_dealer = loaded_dealer
                st.session_state.audit_auditor = loaded_auditor
                st.session_state.audit_date = loaded_date
                st.session_state.selected_claim = next(iter(claims.keys()))
                st.success(f"Auditoría cargada: {len(claims)} claims.")
                st.rerun()
            except Exception as exc:
                st.error(f"No se pudo cargar el JSON: {exc}")

        st.divider()
        st.subheader("Regla de puntuación")
        st.write(f"Documentación: **{MAX_DOCUMENT_POINTS}**")
        st.write(f"Piezas viejas: **{MAX_OLD_PARTS_POINTS}**")
        st.write(f"Total: **{MAX_TOTAL_POINTS}**")
        st.caption("Por defecto: OK / máximo. No aplica = máximo del apartado.")

    claims: Dict[str, Dict[str, Any]] = st.session_state.claims
    apply_default_dealer_to_blank_claims(claims, safe_str(st.session_state.audit_dealer))
    sync_all_claims_from_widget_state(claims)

    if not claims:
        st.info("Pega la lista de claims, con dealer si aplica, en la tabla de abajo para empezar.")
        render_claim_paste_loader(dealer, show_title=True)
        st.stop()

    with st.expander("Añadir o reemplazar claims desde tabla", expanded=False):
        render_claim_paste_loader(dealer, show_title=False)

    audit_score = calculate_audit_score(claims)
    kpi_cols = st.columns(5)
    kpi_cols[0].metric("Claims", audit_score["claims"])
    kpi_cols[1].metric("Completadas", audit_score["completed_claims"])
    kpi_cols[2].metric("Documentación", f"{audit_score['doc_points']}/{audit_score['claims'] * MAX_DOCUMENT_POINTS}")
    kpi_cols[3].metric("Piezas viejas", f"{audit_score['old_points']}/{audit_score['claims'] * MAX_OLD_PARTS_POINTS}")
    kpi_cols[4].metric("Éxito global", f"{audit_score['success_percent']:.1f}%")
    st.divider()

    left, right = st.columns([1.05, 2.5])

    with left:
        st.subheader("Claims")
        summary_df = build_summary_dataframe(claims, "es")
        compact_cols = [TEXT["es"]["claim_no"], TEXT["es"]["other_id"], TEXT["es"]["dealer"], TEXT["es"]["deduction_amount"], TEXT["es"]["doc_score"], TEXT["es"]["old_score"], TEXT["es"]["total_score"], TEXT["es"]["result"]]
        st.dataframe(summary_df[[col for col in compact_cols if col in summary_df.columns]], use_container_width=True, hide_index=True)

        claim_options = list(claims.keys())
        if st.session_state.selected_claim not in claim_options:
            st.session_state.selected_claim = claim_options[0]
        selected_claim = st.selectbox("Seleccionar claim", claim_options, index=claim_options.index(st.session_state.selected_claim), format_func=display_claim_option)
        st.session_state.selected_claim = selected_claim

        st.download_button("💾 Guardar progreso editable (.json)", data=serialize_audit_workfile(claims, audit_name, dealer, auditor, audit_date_value), file_name=f"{base_name}.json", mime="application/json")
        st.download_button("🇪🇸 Boletín ES para dealer", data=export_scorecard_excel(claims, audit_name, dealer, auditor, "es", audit_date_value), file_name=f"{base_name}_ES_boletin.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        st.download_button("🇬🇧 Scorecard EN para HQ", data=export_scorecard_excel(claims, audit_name, dealer, auditor, "en", audit_date_value), file_name=f"{base_name}_EN_scorecard.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        st.download_button("Exportar analítico ES", data=export_analytical_excel(claims, audit_name, dealer, auditor, "es", audit_date_value), file_name=f"{base_name}_ES_analitico.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        st.download_button("Exportar paquete ES+EN (.zip)", data=export_bilingual_package_zip(claims, audit_name, dealer, auditor, audit_date_value), file_name=f"{base_name}_ES_EN.zip", mime="application/zip")

    with right:
        claim = claims[st.session_state.selected_claim]
        score = calculate_claim_score(claim)
        st.subheader(f"Revisión claim {claim_identifier(claim, 'es') or claim_identifier(claim, 'en')}")
        render_claim_quick_card(claim, dealer)
        render_deduction_editor(claim)

        score_cols = st.columns(4)
        score_cols[0].metric("Documentación", f"{score['doc_points']}/{MAX_DOCUMENT_POINTS}")
        score_cols[1].metric("Piezas viejas", f"{score['old_points']}/{MAX_OLD_PARTS_POINTS}")
        score_cols[2].metric("Resultado claim", f"{score['total_points']}/100")
        score_cols[3].metric("Estado", "Completada" if score["completed"] else "Pendiente")
        if score["pending"]:
            st.warning("Apartados pendientes: " + ", ".join(score["pending"]))

        section_names = ["I. Documentación", "II. Piezas viejas", "Comentarios", "Informe"]
        section_tuple = tuple(section_names)

        # Guardamos el apartado activo por índice propio, no usando la misma key del widget.
        # Así evitamos el error de Streamlit al intentar cambiar una key ya usada por un widget.
        section_index = clamp_index(st.session_state.get("audit_section_index", 0), len(section_names))
        st.session_state.audit_section_index = section_index

        # La key del radio se sincroniza ANTES de crear el widget.
        # Los botones la cambian mediante callback, nunca después de instanciarla.
        current_section_label = section_names[section_index]
        if st.session_state.get("audit_section_radio") != current_section_label:
            st.session_state.audit_section_radio = current_section_label

        selected_section = st.radio(
            "Sección",
            section_names,
            horizontal=True,
            label_visibility="collapsed",
            key="audit_section_radio",
            on_change=sync_audit_section_from_radio,
            args=(section_tuple,),
        )

        section_index = clamp_index(st.session_state.get("audit_section_index", section_names.index(selected_section)), len(section_names))
        selected_section = section_names[section_index]

        if selected_section == "I. Documentación":
            render_check_editor(claim, DOCUMENT_CHECKS)
        elif selected_section == "II. Piezas viejas":
            render_check_editor(claim, OLD_PARTS_CHECKS)
        elif selected_section == "Comentarios":
            render_comments_editor(claim)
        elif selected_section == "Informe":
            render_report_section(claims, audit_name, dealer, auditor, base_name, audit_date_value)

        st.divider()
        nav_section_cols = st.columns([1, 1, 2])
        with nav_section_cols[0]:
            st.button(
                "← Apartado anterior",
                disabled=section_index == 0,
                use_container_width=True,
                on_click=set_audit_section_index,
                args=(section_index - 1, section_tuple),
            )
        with nav_section_cols[1]:
            st.button(
                "Siguiente apartado →",
                type="primary",
                disabled=section_index >= len(section_names) - 1,
                use_container_width=True,
                on_click=set_audit_section_index,
                args=(section_index + 1, section_tuple),
            )
        with nav_section_cols[2]:
            st.caption(f"Apartado {section_index + 1} de {len(section_names)} · {selected_section}")

        st.divider()
        current_index = claim_options.index(st.session_state.selected_claim)
        nav_claim_cols = st.columns([1, 1, 2])
        with nav_claim_cols[0]:
            st.button(
                "← Anterior claim",
                disabled=current_index == 0,
                use_container_width=True,
                on_click=set_selected_claim_keep_section,
                args=(claim_options[current_index - 1] if current_index > 0 else st.session_state.selected_claim, section_tuple),
            )
        with nav_claim_cols[1]:
            st.button(
                "Siguiente claim →",
                disabled=current_index >= len(claim_options) - 1,
                use_container_width=True,
                on_click=set_selected_claim_keep_section,
                args=(claim_options[current_index + 1] if current_index < len(claim_options) - 1 else st.session_state.selected_claim, section_tuple),
            )
        with nav_claim_cols[2]:
            st.caption(
                f"Claim {current_index + 1} de {len(claim_options)} · "
                "La navegación mantiene el apartado activo para poder revisar todas las claims por bloque."
            )

        # Accesos rápidos para el flujo habitual:
        # 1) revisar toda la documentación claim a claim;
        # 2) ir al almacén y revisar piezas viejas claim a claim.
        if section_index == 0 and current_index == len(claim_options) - 1 and len(section_names) > 1:
            st.button(
                "He terminado documentación → empezar II. Piezas viejas",
                type="primary",
                use_container_width=True,
                on_click=set_selected_claim_and_section,
                args=(claim_options[0], 1, section_tuple),
                key="start_old_parts_after_documents",
            )
        elif section_index == 1 and current_index == 0:
            st.caption("Estás en II. Piezas viejas. El botón 'Siguiente claim' seguirá en piezas viejas para trabajar desde almacén.")


if __name__ == "__main__":
    main()
